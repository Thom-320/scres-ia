"""
Final polish pass on all new tables (indices 3-5, 7-9) in the updated v0.
Ensures:
- Explicit column widths (proportional, totaling page width minus margins)
- Consistent cell padding
- Header row shading (light blue D9E2F3)
- All borders uniform
- Text alignment: headers centered, data left-aligned for text, center for numbers
- Font: Times New Roman 11pt in tables, bold headers
- Row heights reasonable
- No auto-fit (fixed widths)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = 'Times New Roman'
FONT_SIZE_TABLE = Pt(10)  # Slightly smaller for tables = cleaner
FONT_SIZE_HEADER = Pt(10)

# Page width minus margins = usable width
# Page: 7773670 EMU, margins: 900430 each side
USABLE_WIDTH_EMU = 7773670 - 2 * 900430  # = 5972810 EMU
USABLE_WIDTH_CM = USABLE_WIDTH_EMU / 360000  # ~16.6 cm

doc_path = os.path.expanduser('~/Downloads/v0_neuralNet-scres_UPDATED_2026-03-24.docx')
doc = Document(doc_path)


def set_cell_width(cell, width_emu):
    """Set explicit cell width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="0" w:type="dxa"/>')
        tcPr.append(tcW)
    # Convert EMU to twips (1 twip = 1/1440 inch, 1 EMU = 1/914400 inch)
    twips = int(width_emu / 635)
    tcW.set(qn('w:w'), str(twips))
    tcW.set(qn('w:type'), 'dxa')


def set_cell_margins(cell, top=40, bottom=40, left=60, right=60):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    # Remove existing margins if any
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(margins)


def set_cell_shading(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    old = tcPr.find(qn('w:shd'))
    if old is not None:
        tcPr.remove(old)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)


def format_cell_text(cell, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=FONT_SIZE_TABLE):
    """Reformat all text in a cell."""
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = font_size
            run.bold = bold
            run.font.color.rgb = None  # Reset to auto/black


def set_table_borders(table, color="000000", size="6"):
    """Set clean uniform borders on a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    # Remove existing borders
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_layout_fixed(table):
    """Set table to fixed layout (no auto-fit)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    old_layout = tblPr.find(qn('w:tblLayout'))
    if old_layout is not None:
        tblPr.remove(old_layout)
    
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tblPr.append(layout)


def set_table_width(table, width_twips):
    """Set explicit table width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    tblPr.append(tblW)


def polish_table(table, col_widths_pct, header_color="D9E2F3", text_aligns=None):
    """
    Polish a table with:
    - Fixed column widths (percentages of usable width)
    - Uniform borders
    - Header shading
    - Consistent fonts
    - Cell margins
    """
    total_twips = int(USABLE_WIDTH_EMU / 635)
    
    # Set table properties
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    set_table_width(table, total_twips)
    set_table_borders(table, color="000000", size="4")  # Clean black borders
    
    col_widths_twips = [int(total_twips * pct) for pct in col_widths_pct]
    
    for row_idx, row in enumerate(table.rows):
        # Set row height
        row.height = Pt(18) if row_idx == 0 else None
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST if row_idx == 0 else WD_ROW_HEIGHT_RULE.AUTO
        
        for col_idx, cell in enumerate(row.cells):
            # Set width
            set_cell_width(cell, col_widths_twips[col_idx] * 635)
            
            # Set margins (compact)
            set_cell_margins(cell, top=30, bottom=30, left=50, right=50)
            
            # Header row: shaded, bold, centered
            if row_idx == 0:
                set_cell_shading(cell, header_color)
                format_cell_text(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font_size=FONT_SIZE_HEADER)
            else:
                # Data rows: clean white, no zebra striping
                set_cell_shading(cell, "FFFFFF")
                
                # Alignment per column
                if text_aligns and col_idx < len(text_aligns):
                    align = text_aligns[col_idx]
                else:
                    align = WD_ALIGN_PARAGRAPH.CENTER
                
                format_cell_text(cell, bold=False, align=align, font_size=FONT_SIZE_TABLE)


# =====================================================================
# POLISH EACH NEW TABLE
# =====================================================================

# Table 3: Observation vector (16 rows x 4 cols) - Table 5 in paper
print("Polishing Table 3 (observation vector)...")
polish_table(
    doc.tables[3],
    col_widths_pct=[0.07, 0.22, 0.45, 0.26],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,  # Index
        WD_ALIGN_PARAGRAPH.LEFT,     # Variable  
        WD_ALIGN_PARAGRAPH.LEFT,     # Description
        WD_ALIGN_PARAGRAPH.CENTER,   # Normalization
    ]
)

# Table 4: Action space (6 rows x 4 cols) - Table 6 in paper
print("Polishing Table 4 (action space)...")
polish_table(
    doc.tables[4],
    col_widths_pct=[0.06, 0.14, 0.48, 0.32],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,  # Dim
        WD_ALIGN_PARAGRAPH.LEFT,     # Action
        WD_ALIGN_PARAGRAPH.LEFT,     # Mapping
        WD_ALIGN_PARAGRAPH.LEFT,     # DES parameter
    ]
)

# Table 5: PPO hyperparameters (10 rows x 2 cols) - Table 7 in paper
print("Polishing Table 5 (PPO hyperparameters)...")
polish_table(
    doc.tables[5],
    col_widths_pct=[0.55, 0.45],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Parameter
        WD_ALIGN_PARAGRAPH.CENTER,   # Value
    ]
)

# Table 7: Static baselines (7 rows x 5 cols) - Table 8 in paper
print("Polishing Table 7 (static baselines)...")
polish_table(
    doc.tables[7],
    col_widths_pct=[0.15, 0.17, 0.24, 0.22, 0.22],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,  # Scenario
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.CENTER,   # Control reward
        WD_ALIGN_PARAGRAPH.CENTER,   # Fill rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Backorder rate
    ]
)

# Table 8: PPO vs static increased (5 rows x 5 cols) - Table 9 in paper
print("Polishing Table 8 (PPO increased risk)...")
polish_table(
    doc.tables[8],
    col_widths_pct=[0.16, 0.19, 0.15, 0.19, 0.31],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.CENTER,   # Control reward
        WD_ALIGN_PARAGRAPH.CENTER,   # Fill rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ]
)

# Table 9: PPO vs static severe (5 rows x 5 cols) - Table 10 in paper
print("Polishing Table 9 (PPO severe risk)...")
polish_table(
    doc.tables[9],
    col_widths_pct=[0.16, 0.19, 0.15, 0.19, 0.31],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.CENTER,   # Control reward
        WD_ALIGN_PARAGRAPH.CENTER,   # Fill rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ]
)

# Save
output_path = os.path.expanduser('~/Downloads/v0_neuralNet-scres_UPDATED_2026-03-24.docx')
doc.save(output_path)
print(f"\nSaved polished version to: {output_path}")

# Also regenerate PDF
import subprocess
subprocess.run([
    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    '--headless', '--convert-to', 'pdf',
    '--outdir', os.path.expanduser('~/Downloads'),
    output_path
], capture_output=True)
print("PDF regenerated.")

"""
Rebuild all new tables from scratch at Q1 journal standard.
Target: IJPR / EJOR / IEEE TAI

Key requirements:
- Booktabs: toprule (sz=18), midrule (sz=6), bottomrule (sz=18)
- NO vertical lines anywhere
- NO shading anywhere (not even headers)
- NO table style (remove tblStyle to prevent Word from adding shading)
- Right-aligned numerics
- Bold best results
- Grouping midrules for long tables
- Compact rows
- Times New Roman 10pt
"""

import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(10)
USABLE_EMU = 7773670 - 2 * 900430
TOTAL_TWIPS = int(USABLE_EMU / 635)

# Booktabs rule sizes (in half-points)
TOPRULE = "24"     # Thick (3pt)
MIDRULE = "8"      # Thin (1pt)
BOTTOMRULE = "24"  # Thick (3pt)
CMIDRULE = "6"     # Thin (0.75pt, for grouping — same as midrule but grey)


def nuke_table_style(table):
    """Remove table style entirely so Word doesn't inject shading."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is not None:
        for child in list(tblPr):
            tag = etree.QName(child.tag).localname
            if tag in ('tblStyle', 'tblLook'):
                tblPr.remove(child)


def set_no_borders_table_level(table):
    """Set all table-level borders to none."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblLayout'))
    if old is not None:
        tblPr.remove(old)
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tblPr.append(layout)
    # Width
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{TOTAL_TWIPS}" w:type="dxa"/>')
    tblPr.append(tblW)


def set_cell_borders(cell, top=None, bottom=None):
    """Set specific borders on a cell. top/bottom = (val, sz, color) or None."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    
    parts = []
    if top:
        parts.append(f'<w:top w:val="{top[0]}" w:sz="{top[1]}" w:space="0" w:color="{top[2]}"/>')
    if bottom:
        parts.append(f'<w:bottom w:val="{bottom[0]}" w:sz="{bottom[1]}" w:space="0" w:color="{bottom[2]}"/>')
    
    if parts:
        xml = f'<w:tcBorders {nsdecls("w")}>{"".join(parts)}</w:tcBorders>'
        tcPr.append(parse_xml(xml))


def clean_cell(cell, width_twips, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Format a cell completely: width, margins, no shading, font."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Width
    old_w = tcPr.find(qn('w:tcW'))
    if old_w is not None:
        tcPr.remove(old_w)
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    tcPr.append(tcW)
    
    # Margins (compact)
    old_mar = tcPr.find(qn('w:tcMar'))
    if old_mar is not None:
        tcPr.remove(old_mar)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        '  <w:top w:w="30" w:type="dxa"/>'
        '  <w:left w:w="50" w:type="dxa"/>'
        '  <w:bottom w:w="30" w:type="dxa"/>'
        '  <w:right w:w="50" w:type="dxa"/>'
        '</w:tcMar>'
    )
    tcPr.append(margins)
    
    # Remove ALL shading
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    
    # Font and alignment
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        pPr = paragraph._element.find(qn('w:pPr'))
        if pPr is not None:
            for shd in pPr.findall(qn('w:shd')):
                pPr.remove(shd)
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE
            run.bold = bold
            run.font.color.rgb = None
            rPr = run._element.find(qn('w:rPr'))
            if rPr is not None:
                for shd in rPr.findall(qn('w:shd')):
                    rPr.remove(shd)


def booktabs(table, col_pcts, aligns, bold_cells=None, group_after_rows=None):
    """
    Full booktabs rebuild.
    
    col_pcts: list of column width fractions
    aligns: list of WD_ALIGN_PARAGRAPH per column
    bold_cells: set of (row, col) to bold
    group_after_rows: list of row indices after which to add a thin cmidrule
    """
    nuke_table_style(table)
    set_no_borders_table_level(table)
    set_table_fixed(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_twips = [int(TOTAL_TWIPS * p) for p in col_pcts]
    n_rows = len(table.rows)
    group_rows = set(group_after_rows or [])
    
    for ri, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        
        for ci, cell in enumerate(row.cells):
            is_header = (ri == 0)
            is_bold = is_header or (bold_cells and (ri, ci) in bold_cells)
            align = WD_ALIGN_PARAGRAPH.CENTER if is_header else aligns[ci]
            
            clean_cell(cell, col_twips[ci], bold=is_bold, align=align)
            
            # Booktabs rules
            top_border = None
            bottom_border = None
            
            # Header row: toprule on top, midrule on bottom
            if ri == 0:
                top_border = ("single", TOPRULE, "000000")
                bottom_border = ("single", MIDRULE, "000000")
            
            # Last row: bottomrule on bottom
            if ri == n_rows - 1:
                bottom_border = ("single", BOTTOMRULE, "000000")
            
            # Group separator: thin cmidrule
            if ri in group_rows:
                bottom_border = ("single", CMIDRULE, "999999")
            
            # First data row after a group: add thin top
            if ri - 1 in group_rows and ri > 0:
                # Already handled by the group_row's bottom border
                pass
            
            set_cell_borders(cell, top=top_border, bottom=bottom_border)


# =====================================================================
# Load fresh base
# =====================================================================

# First regenerate the base document
base_dir = os.path.expanduser('~/Desktop/Universidad_Codigo/proyecto_grarrido_scres+ia')
subprocess.run(['python3', 'scripts/edit_v0_inplace.py'], cwd=base_dir, capture_output=True)

doc_path = os.path.expanduser('~/Downloads/v0_neuralNet-scres_UPDATED_2026-03-24.docx')
doc = Document(doc_path)

# =====================================================================
# Table 3 (paper Table 5): Observation vector — 16 rows x 4 cols
# Group: inventory (rows 1-6), rates (7-8), binary (9-12), meta (13-15)
# =====================================================================
print("Rebuilding Table 5 (observation vector)...")
booktabs(
    doc.tables[3],
    col_pcts=[0.07, 0.22, 0.45, 0.26],
    aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,   # Index
        WD_ALIGN_PARAGRAPH.LEFT,     # Variable
        WD_ALIGN_PARAGRAPH.LEFT,     # Description
        WD_ALIGN_PARAGRAPH.CENTER,   # Normalization
    ],
    group_after_rows=[6, 12],  # After inventory block (0-6), after binary indicators (8-12)
)

# =====================================================================
# Table 4 (paper Table 6): Action space — 6 rows x 4 cols
# =====================================================================
print("Rebuilding Table 6 (action space)...")
booktabs(
    doc.tables[4],
    col_pcts=[0.06, 0.14, 0.48, 0.32],
    aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,   # Dim
        WD_ALIGN_PARAGRAPH.LEFT,     # Action
        WD_ALIGN_PARAGRAPH.LEFT,     # Mapping
        WD_ALIGN_PARAGRAPH.LEFT,     # DES parameter
    ],
    group_after_rows=[4],  # Separate inventory actions from shift action
)

# =====================================================================
# Table 5 (paper Table 7): PPO hyperparameters — 10 rows x 2 cols
# =====================================================================
print("Rebuilding Table 7 (PPO hyperparameters)...")
booktabs(
    doc.tables[5],
    col_pcts=[0.55, 0.45],
    aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Parameter
        WD_ALIGN_PARAGRAPH.RIGHT,    # Value
    ],
)

# =====================================================================
# Table 7 (paper Table 8): Static baselines — 7 rows x 5 cols
# Bold best: S2 under increased (row 2), S2 under severe (row 5)
# Group: separate increased from severe
# =====================================================================
print("Rebuilding Table 8 (static baselines)...")
booktabs(
    doc.tables[7],
    col_pcts=[0.15, 0.17, 0.24, 0.22, 0.22],
    aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Scenario
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
    ],
    bold_cells={
        (2, 2), (2, 3), (2, 4),  # S2 increased = best reward, fill, backorder
        (5, 2),                    # S2 severe = best reward
        (6, 3), (6, 4),           # S3 severe = best fill, backorder
    },
    group_after_rows=[3],  # Separate increased from severe
)

# =====================================================================
# Table 8 (paper Table 9): PPO increased — 5 rows x 5 cols
# =====================================================================
print("Rebuilding Table 9 (PPO increased risk)...")
booktabs(
    doc.tables[8],
    col_pcts=[0.18, 0.19, 0.13, 0.19, 0.31],
    aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ],
    group_after_rows=[2],  # Separate policies from difference/CI
)

# =====================================================================
# Table 9 (paper Table 10): PPO severe — 5 rows x 5 cols
# Bold PPO reward (better)
# =====================================================================
print("Rebuilding Table 10 (PPO severe risk)...")
booktabs(
    doc.tables[9],
    col_pcts=[0.18, 0.19, 0.13, 0.19, 0.31],
    aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ],
    bold_cells={
        (2, 1),  # PPO reward = best (−380.98 > −385.59)
    },
    group_after_rows=[2],  # Separate policies from difference/CI
)

# =====================================================================
# Save
# =====================================================================
doc.save(doc_path)
print(f"\nSaved Q1-grade version to: {doc_path}")

# Regenerate PDF
subprocess.run([
    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    '--headless', '--convert-to', 'pdf',
    '--outdir', os.path.expanduser('~/Downloads'),
    doc_path
], capture_output=True)
print("PDF regenerated.")

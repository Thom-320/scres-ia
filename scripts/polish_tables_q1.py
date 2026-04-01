"""
Q1 journal-grade table polish (IJPR/EJOR/IEEE TAI standard).

Booktabs style:
- NO vertical borders
- Only 3 horizontal rules: top (thick), header separator (thin), bottom (thick)
- No zebra striping, no cell shading (not even headers)
- Right-aligned numeric data
- Bold best results in comparison tables
- Compact row heights
- Times New Roman 10pt in tables
- Table captions: "Table X." bold + description in regular
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FONT_NAME = 'Times New Roman'
FONT_SIZE_TABLE = Pt(10)

USABLE_WIDTH_EMU = 7773670 - 2 * 900430
TOTAL_TWIPS = int(USABLE_WIDTH_EMU / 635)


def set_cell_width(cell, width_twips):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="0" w:type="dxa"/>')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_twips))
    tcW.set(qn('w:type'), 'dxa')


def set_cell_margins(cell, top=20, bottom=20, left=40, right=40):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcMar'))
    if old is not None:
        tcPr.remove(old)
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)


def clear_cell_shading(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:shd'))
    if old is not None:
        tcPr.remove(old)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="auto" w:val="clear"/>')
    tcPr.append(shading)


def format_cell_text(cell, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(1)
        paragraph.paragraph_format.space_after = Pt(1)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_TABLE
            run.bold = bold
            run.font.color.rgb = None


def set_cell_bottom_border(cell, sz="6", color="000000"):
    """Add bottom border to a specific cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        tcPr.remove(old)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)


def set_cell_top_border(cell, sz="12", color="000000"):
    """Add top border to a specific cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    old = tcPr.find(qn('w:tcBorders'))
    if old is not None:
        # Preserve existing borders, add top
        existing_bottom = old.find(qn('w:bottom'))
        bottom_attrs = {}
        if existing_bottom is not None:
            bottom_attrs = dict(existing_bottom.attrib)
        tcPr.remove(old)
        borders_xml = f'<w:tcBorders {nsdecls("w")}>'
        borders_xml += f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        if bottom_attrs:
            b_sz = bottom_attrs.get(qn('w:sz'), '6')
            b_color = bottom_attrs.get(qn('w:color'), '000000')
            borders_xml += f'  <w:bottom w:val="single" w:sz="{b_sz}" w:space="0" w:color="{b_color}"/>'
        borders_xml += '</w:tcBorders>'
        borders = parse_xml(borders_xml)
        tcPr.append(borders)
    else:
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)


def remove_table_borders(table):
    """Remove all table-level borders (we'll add cell-level borders for booktabs)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    # Set all borders to none at table level
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_layout_fixed(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblLayout'))
    if old is not None:
        tblPr.remove(old)
    layout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
    tblPr.append(layout)


def set_table_width(table, width_twips):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblW'))
    if old is not None:
        tblPr.remove(old)
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{width_twips}" w:type="dxa"/>')
    tblPr.append(tblW)


def booktabs_polish(table, col_widths_pct, text_aligns=None, bold_cells=None):
    """
    Apply booktabs-style formatting:
    - Thick top rule on first row
    - Thin rule below header
    - Thick bottom rule on last row
    - No vertical lines
    - No shading
    - Compact cells
    """
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_layout_fixed(table)
    set_table_width(table, TOTAL_TWIPS)
    remove_table_borders(table)
    
    col_widths = [int(TOTAL_TWIPS * pct) for pct in col_widths_pct]
    n_rows = len(table.rows)
    
    for row_idx, row in enumerate(table.rows):
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[col_idx])
            set_cell_margins(cell, top=25, bottom=25, left=40, right=40)
            clear_cell_shading(cell)
            
            # Alignment
            is_header = (row_idx == 0)
            if is_header:
                align = WD_ALIGN_PARAGRAPH.CENTER
            elif text_aligns and col_idx < len(text_aligns):
                align = text_aligns[col_idx]
            else:
                align = WD_ALIGN_PARAGRAPH.CENTER
            
            # Bold for headers or specific cells
            cell_bold = is_header
            if bold_cells and (row_idx, col_idx) in bold_cells:
                cell_bold = True
            
            format_cell_text(cell, bold=cell_bold, align=align)
            
            # Booktabs rules:
            # Top thick rule on header row
            if row_idx == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                old_b = tcPr.find(qn('w:tcBorders'))
                if old_b is not None:
                    tcPr.remove(old_b)
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                    '  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(borders)
            
            # Bottom thick rule on last row
            elif row_idx == n_rows - 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                old_b = tcPr.find(qn('w:tcBorders'))
                if old_b is not None:
                    tcPr.remove(old_b)
                borders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                    '</w:tcBorders>'
                )
                tcPr.append(borders)
            
            # All other rows: no borders
            else:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                old_b = tcPr.find(qn('w:tcBorders'))
                if old_b is not None:
                    tcPr.remove(old_b)


# =====================================================================
# Load and polish
# =====================================================================

doc_path = os.path.expanduser('~/Downloads/v0_neuralNet-scres_UPDATED_2026-03-24.docx')
doc = Document(doc_path)

# Table 3: Observation vector (Table 5 in paper) - 16 rows x 4 cols
print("Polishing Table 5 (observation vector)...")
booktabs_polish(
    doc.tables[3],
    col_widths_pct=[0.07, 0.22, 0.45, 0.26],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,   # Index
        WD_ALIGN_PARAGRAPH.LEFT,     # Variable  
        WD_ALIGN_PARAGRAPH.LEFT,     # Description
        WD_ALIGN_PARAGRAPH.CENTER,   # Normalization
    ]
)

# Table 4: Action space (Table 6) - 6 rows x 4 cols
print("Polishing Table 6 (action space)...")
booktabs_polish(
    doc.tables[4],
    col_widths_pct=[0.06, 0.14, 0.48, 0.32],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.CENTER,   # Dim
        WD_ALIGN_PARAGRAPH.LEFT,     # Action
        WD_ALIGN_PARAGRAPH.LEFT,     # Mapping
        WD_ALIGN_PARAGRAPH.LEFT,     # DES parameter
    ]
)

# Table 5: PPO hyperparameters (Table 7) - 10 rows x 2 cols
print("Polishing Table 7 (PPO hyperparameters)...")
booktabs_polish(
    doc.tables[5],
    col_widths_pct=[0.55, 0.45],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Parameter
        WD_ALIGN_PARAGRAPH.RIGHT,    # Value (right-aligned for numbers)
    ]
)

# Table 7: Static baselines (Table 8) - 7 rows x 5 cols
# Bold best results: S2 under increased (row 2), S2 under severe (row 5)
print("Polishing Table 8 (static baselines)...")
booktabs_polish(
    doc.tables[7],
    col_widths_pct=[0.15, 0.17, 0.24, 0.22, 0.22],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Scenario
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
    ],
    bold_cells={
        (2, 2), (2, 3), (2, 4),  # Static S2 increased = best
        (5, 2), (6, 3), (6, 4),  # S2 best reward severe, S3 best fill severe
    }
)

# Table 8: PPO vs static increased (Table 9) - 5 rows x 5 cols
# Bold PPO fill rate (marginally better)
print("Polishing Table 9 (PPO increased risk)...")
booktabs_polish(
    doc.tables[8],
    col_widths_pct=[0.18, 0.19, 0.13, 0.19, 0.31],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ]
)

# Table 9: PPO vs static severe (Table 10) - 5 rows x 5 cols
# Bold PPO reward (better)
print("Polishing Table 10 (PPO severe risk)...")
booktabs_polish(
    doc.tables[9],
    col_widths_pct=[0.18, 0.19, 0.13, 0.19, 0.31],
    text_aligns=[
        WD_ALIGN_PARAGRAPH.LEFT,     # Policy
        WD_ALIGN_PARAGRAPH.RIGHT,    # Control reward
        WD_ALIGN_PARAGRAPH.RIGHT,    # Fill rate
        WD_ALIGN_PARAGRAPH.RIGHT,    # Backorder rate
        WD_ALIGN_PARAGRAPH.CENTER,   # Shift mix
    ],
    bold_cells={
        (1, 1),  # PPO reward (best)
    }
)

# Save
doc.save(doc_path)
print(f"\nSaved Q1-polished version to: {doc_path}")

# Regenerate PDF
import subprocess
subprocess.run([
    '/Applications/LibreOffice.app/Contents/MacOS/soffice',
    '--headless', '--convert-to', 'pdf',
    '--outdir', os.path.expanduser('~/Downloads'),
    doc_path
], capture_output=True)
print("PDF regenerated.")

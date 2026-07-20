#!/usr/bin/env python3
"""Build the Paper 2 pre-results manuscript DOCX from its authoritative Markdown."""

from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript/paper2/PAPER2_MANUSCRIPT_DRAFT.md"
OUTPUT = ROOT / "manuscript/paper2/PAPER2_MANUSCRIPT_DRAFT.docx"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)


def set_font(run, *, name: str = "Calibri", size: float = 11, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, end))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_inline(paragraph, text: str) -> None:
    # Minimal Markdown emphasis/code renderer; equations remain literal and
    # auditable instead of being converted to images.
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            set_font(paragraph.add_run(text[cursor : match.start()]))
        token = match.group(0)
        if token.startswith("`"):
            set_font(paragraph.add_run(token[1:-1]), name="Courier New", size=9.5)
        else:
            set_font(paragraph.add_run(token[2:-2]), bold=True)
        cursor = match.end()
    if cursor < len(text):
        set_font(paragraph.add_run(text[cursor:]))


def build() -> None:
    if OUTPUT.exists():
        OUTPUT.unlink()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("PAPER 2 | PRE-RESULTS MANUSCRIPT"), size=8.5, bold=True, color=GRAY)
    page_field(section.footer.paragraphs[0])

    lines = SOURCE.read_text().splitlines()
    first = True
    paragraph_buffer: list[str] = []

    def flush() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, " ".join(paragraph_buffer))
            paragraph_buffer = []

    for raw in lines:
        line = raw.strip()
        if not line:
            flush()
            continue
        if line.startswith("# ") and first:
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(36)
            p.paragraph_format.space_after = Pt(10)
            set_font(p.add_run(line[2:]), size=23, bold=True, color=DARK_BLUE)
            first = False
        elif line.startswith("## "):
            flush()
            text = line[3:]
            if text == "Exact static benchmarks and full discrete-event validation":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(18)
                set_font(p.add_run(text), size=13, italic=True, color=GRAY)
            else:
                doc.add_paragraph(text, style="Heading 1")
        elif line.startswith("### "):
            flush()
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("> "):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.right_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            set_font(p.add_run(line[2:]), size=10.5, italic=True, color=DARK_BLUE)
        elif re.match(r"^\d+\. ", line):
            flush()
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.194)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            flush()
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.375)
            p.paragraph_format.first_line_indent = Inches(-0.194)
            p.paragraph_format.space_after = Pt(4)
            add_inline(p, line[2:])
        elif line.startswith("\\[") or line == "\\]":
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_font(p.add_run(line), name="Cambria Math", size=10.5)
        elif line.startswith("**Target journal:**") or line.startswith("**Document status:**") or line.startswith("**Model status:**"):
            flush()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, line.rstrip("  "))
        else:
            paragraph_buffer.append(line.rstrip("  "))
    flush()
    core = doc.core_properties
    core.title = "Learning-Augmented Event-Triggered MPC for Supply-Chain Resilience"
    core.subject = "Pre-results manuscript for Computers & Industrial Engineering"
    core.keywords = "SCRES, DES, MPC, reinforcement learning, event-triggered control"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)


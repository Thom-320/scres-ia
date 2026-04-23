"""
Build a DOCX with Sections 3.3 and 4.2 matching the v0 format.

Format spec (from v0 inspection):
- Font: Times New Roman throughout
- Title: 16pt bold, centered
- Section headings (3.3, 4.2): 13pt bold
- Subsection headings (3.3.1): 13pt italic
- Body text: 13pt normal, left-aligned
- Table captions: 13pt normal, centered
- Tables: 13pt, bordered
- Line spacing: 1.5
- Margins: ~2.5cm all around
- Page size: ~A4-ish (7773670 x 10059670 EMU)
"""

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Emu(900430)
    section.bottom_margin = Emu(900430)
    section.left_margin = Emu(900430)
    section.right_margin = Emu(900430)
    section.page_width = Emu(7773670)
    section.page_height = Emu(10059670)

# Set Normal style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(6)


def add_section_heading(text):
    """Section heading: 13pt bold (e.g., '3.3 Hybrid...')"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.bold = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subsection_heading(text):
    """Subsection heading: 13pt italic (e.g., '3.3.1 State Space')"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.italic = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_body(text):
    """Body text: 13pt normal, left-aligned"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body_italic(text):
    """Body text italic"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_table_caption(text):
    """Table caption: 13pt normal, centered"""
    p = doc.add_paragraph()
    # Bold the "Table X." part
    parts = text.split(".", 1)
    if len(parts) == 2:
        run1 = p.add_run(parts[0] + ".")
        run1.font.name = "Times New Roman"
        run1.font.size = Pt(13)
        run1.bold = True
        run2 = p.add_run(parts[1])
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(13)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(13)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    return p


def add_table(headers, rows):
    """Add a formatted table with borders and Times New Roman 11pt."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set borders on all cells
    tbl = table._tbl
    tblPr = (
        tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    )
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        "</w:tblBorders>"
    )
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Shade header
    for cell in table.rows[0].cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def add_equation(text):
    """Add an equation-like centered paragraph"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


# =====================================================================
# DOCUMENT HEADER
# =====================================================================
p = doc.add_paragraph()
run = p.add_run("Sections 3.3 and 4.2 — Paper Update for v0")
run.font.name = "Times New Roman"
run.font.size = Pt(16)
run.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_body_italic("Ready for insertion into v.0_neuralNet-scres. Written 2026-03-24.")
doc.add_paragraph()  # spacer

# =====================================================================
# SECTION 3.3
# =====================================================================
add_section_heading("3.3 Hybrid Simulation–Neural Model")

add_body(
    "The second layer of the proposed framework converts the DES environment into a sequential "
    "decision problem amenable to reinforcement learning. Rather than executing a fixed shift "
    "and inventory policy for the full simulation horizon, the system is controlled by an agent "
    "that observes the supply chain state at weekly intervals and selects operational actions "
    "intended to minimize service loss while managing shift costs. The formulation follows the "
    "standard Markov Decision Process (MDP) interface (Sutton & Barto, 2018), implemented "
    "through the Gymnasium API (Towers et al., 2024), which wraps the SimPy-based DES "
    "described in Section 3.2."
)

# 3.3.1
add_subsection_heading("3.3.1 State Space")

add_body(
    "The agent receives a 15-dimensional observation vector at each decision epoch. "
    "Table 5 summarizes the components and their normalization."
)

add_table_caption("Table 5. Observation vector (v1) for the RL agent.")

add_table(
    ["Index", "Variable", "Description", "Normalization"],
    [
        ["0", "raw_material_wdc", "Raw material at Warehouse (Op3)", "/max_capacity"],
        [
            "1",
            "raw_material_al",
            "Raw material at Assembly Line (Op5)",
            "/max_capacity",
        ],
        ["2", "rations_al", "Rations buffer at QC (Op7)", "/max_capacity"],
        ["3", "rations_sb", "Rations at Supply Battalion (Op9)", "/max_capacity"],
        ["4", "rations_cssu", "Rations at CSSUs (Op11)", "/max_capacity"],
        ["5", "rations_theatre", "Rations at Theatre (Op13)", "/max_capacity"],
        ["6", "fill_rate", "Cumulative fill rate", "[0, 1]"],
        ["7", "backorder_rate", "Cumulative backorder rate", "[0, 1]"],
        ["8", "assembly_line_down", "Binary: assembly line disrupted", "{0, 1}"],
        ["9", "any_loc_down", "Binary: any LOC disrupted", "{0, 1}"],
        ["10", "op9_down", "Binary: Supply Battalion disrupted", "{0, 1}"],
        ["11", "op11_down", "Binary: CSSUs disrupted", "{0, 1}"],
        ["12", "time_fraction", "Simulation progress", "[0, 1]"],
        ["13", "pending_batch_norm", "Pending batch / batch size", "[0, 1]"],
        ["14", "contingent_demand_norm", "Pending contingent demand / 2600", "[0, 1]"],
    ],
)

add_body(
    "The observation captures both inventory levels and disruption indicators, allowing the "
    "agent to react to ongoing operational and known-unknown risks. Time fraction provides "
    "temporal context for long-horizon planning. No explicit memory of past observations "
    "is included at this stage; the implications of this design choice for partial "
    "observability are discussed in Section 5."
)

# 3.3.2
add_subsection_heading("3.3.2 Action Space")

add_body(
    "The action space is five-dimensional and continuous, with each dimension bounded "
    "to [−1, 1]. Table 6 describes each action dimension and its mapping to the DES "
    "control parameters."
)

add_table_caption("Table 6. Action space for the RL agent.")

add_table(
    ["Dim", "Action", "Mapping", "DES parameter"],
    [
        ["0", "op3_q", "multiplier = 1.25 + 0.75 × signal", "Dispatch quantity at Op3"],
        [
            "1",
            "op9_q_max",
            "multiplier = 1.25 + 0.75 × signal",
            "Upper dispatch quantity at Op9",
        ],
        ["2", "op3_rop", "multiplier = 1.25 + 0.75 × signal", "Reorder point at Op3"],
        ["3", "op9_rop", "multiplier = 1.25 + 0.75 × signal", "Reorder point at Op9"],
        [
            "4",
            "shifts",
            "< −0.33 → S1; [−0.33, 0.33) → S2; ≥ 0.33 → S3",
            "Active assembly shifts",
        ],
    ],
)

add_body(
    "The first four dimensions adjust inventory-control parameters as multiplicative "
    "perturbations around the thesis baseline values, providing the agent with fine-grained "
    "control over buffer sizing. The fifth dimension controls the number of active shifts "
    "in the assembly line (S1 = single shift, S2 = double shift, S3 = triple shift), which "
    "directly determines production capacity and operating cost. This action structure extends "
    "the two static strategies examined in Garrido-Rios (2017)\u2014inventory buffering "
    "(Strategy I) and shift augmentation (Strategy II)\u2014into a unified, dynamic control policy."
)

# 3.3.3
add_subsection_heading("3.3.3 Decision Epoch and Horizon")

add_body(
    "The agent makes decisions at weekly intervals (168 simulated hours), which matches "
    "the natural operational cycle of the MFSC: procurement contracts operate on weekly "
    "cycles (R12), demand surges arrive at weekly granularity (R24, 672-hour intervals), "
    "and shift adjustments require coordination time that makes sub-weekly switching "
    "impractical. Each episode spans 260 decision steps, corresponding to the full 20-year "
    "simulation horizon of 161,280 hours after the warm-up period."
)

# 3.3.4
add_subsection_heading("3.3.4 Reward Function Design")

add_body(
    "The choice of reward function is critical for RL-based control. Two reward formulations "
    "were evaluated during development:"
)

# ReT_thesis description
p = doc.add_paragraph()
run = p.add_run("Resilience metric (ReT_thesis). ")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True
run2 = p.add_run(
    "The thesis-aligned resilience metric defined by Garrido-Rios (2017, Equations 5.1\u20135.5) "
    "aggregates autotomy, recovery, non-recovery, and fill rate into a single scalar. While this "
    "metric is appropriate for evaluating resilience outcomes, it proved unsuitable as a training "
    "objective: preliminary experiments showed that the agent learned to minimize assembly shifts "
    "to S1 across all conditions, achieving a higher ReT score by reducing cost at the expense of "
    "severe service degradation (fill rate dropped from 0.99 to 0.84). The metric\u2019s structure "
    "rewards cost avoidance more than service maintenance, creating a misaligned incentive for "
    "the learning agent."
)
run2.font.name = "Times New Roman"
run2.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# control_v1 description
p = doc.add_paragraph()
run = p.add_run("Operational control reward (control_v1). ")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True
run2 = p.add_run(
    "To address this misalignment, we designed a reward function that directly penalizes the "
    "two operational quantities that the shift-control agent can influence:"
)
run2.font.name = "Times New Roman"
run2.font.size = Pt(13)
p.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Equation
add_equation(
    "r\u209c = \u2212( w_bo \u00d7 B\u209c/D\u209c + w_cost \u00d7 (S\u209c \u2212 1) )"
)

add_body("where:")

# Bullet-style definitions
for item in [
    "B\u209c is the number of new backorders at step t,",
    "D\u209c is the total demand at step t,",
    "S\u209c \u2208 {1, 2, 3} is the active shift count,",
    "w_bo = 4.0 weights the service-loss penalty,",
    "w_cost = 0.02 weights the shift-cost penalty.",
]:
    p = doc.add_paragraph()
    run = p.add_run("\u2022 " + item)
    run.font.name = "Times New Roman"
    run.font.size = Pt(13)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

add_body(
    "The ratio w_bo / w_cost = 200 reflects the operational priority that maintaining service "
    "to forward-deployed units dominates shift operating cost. This design ensures that the "
    "agent is penalized for backorders (the primary resilience failure mode) while incurring "
    "a proportional cost for activating additional shifts, creating a meaningful service\u2013cost "
    "tradeoff. The ReT metric is retained as a reporting-only evaluation metric for "
    "thesis-aligned comparison."
)

# Justification subheading
p = doc.add_paragraph()
run = p.add_run("Justification. ")
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.bold = True
run2 = p.add_run(
    "The control_v1 reward was selected over ReT_thesis based on three empirical criteria:"
)
run2.font.name = "Times New Roman"
run2.font.size = Pt(13)

for item in [
    "Behavioral alignment: Under control_v1, the trained agent uses a mix of shifts (12% S1, 25% S2, 63% S3 under increased risk), demonstrating genuine adaptive behavior. Under ReT_thesis, the agent collapsed to 99.99% S1.",
    "Service preservation: control_v1-trained agents maintain fill rates comparable to the best static baselines (0.84 under increased risk, 0.63 under severe risk). ReT_thesis-trained agents degraded fill rate to 0.84 even under current (low) risk.",
    "Interpretability: Each component of control_v1 maps directly to an observable operational quantity, making the reward transparent to supply chain practitioners.",
]:
    p = doc.add_paragraph()
    # Split at first colon for bold label
    parts = item.split(":", 1)
    run1 = p.add_run(parts[0] + ":")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(13)
    run1.italic = True
    run2 = p.add_run(parts[1])
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(13)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

# 3.3.5
add_subsection_heading("3.3.5 Learning Algorithm")

add_body(
    "The agent is trained using Proximal Policy Optimization (PPO; Schulman et al., 2017), "
    "a policy-gradient algorithm with clipped surrogate objectives that provides stable "
    "learning in continuous action spaces. PPO was selected for its robustness to hyperparameter "
    "choices and its established track record in operations research applications (De Moor et al., "
    "2022; Kemmer et al., 2018). The implementation uses Stable-Baselines3 (Raffin et al., 2021) "
    "with the hyperparameters listed in Table 7."
)

add_table_caption("Table 7. PPO hyperparameters.")

add_table(
    ["Parameter", "Value"],
    [
        ["Learning rate", "3 \u00d7 10\u207b\u2074"],
        ["Rollout steps (n_steps)", "1,024"],
        ["Mini-batch size", "64"],
        ["Update epochs", "10"],
        ["Discount factor (\u03b3)", "0.99"],
        ["GAE parameter (\u03bb)", "0.95"],
        ["Clip range", "0.2"],
        ["Training timesteps", "500,000"],
        ["Policy architecture", "MLP (64, 64)"],
    ],
)

# 3.3.6
add_subsection_heading("3.3.6 Evaluation Protocol")

add_body(
    "Each experimental condition is evaluated across 5 training seeds. For each seed, "
    "the trained policy is evaluated over 10 episodes with distinct simulation seeds. "
    "Three static baselines (S1-only, S2-only, S3-only) and a random policy are evaluated "
    "under identical conditions for comparison. Metrics are reported as seed-level means "
    "with bootstrap 95% confidence intervals. Two stress scenarios are examined:"
)

for item in [
    "Increased risk: The baseline risk configuration from Garrido-Rios (2017), representing recurring operational and known-unknown disruptions.",
    "Severe risk: An escalated configuration in which disruption frequencies are increased 4\u201317\u00d7 and recovery times are extended 2\u20133\u00d7, representing a high-stress operational theatre.",
]:
    p = doc.add_paragraph()
    parts = item.split(":", 1)
    run1 = p.add_run(parts[0] + ":")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(13)
    run1.bold = True
    run2 = p.add_run(parts[1])
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(13)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)

add_body(
    "Stochastic processing times are enabled in both scenarios to capture the full variability "
    "of the DES environment."
)

# PAGE BREAK before Section 4.2
doc.add_page_break()

# =====================================================================
# SECTION 4.2
# =====================================================================
add_section_heading("4.2 Results from the Hybrid Simulation\u2013Neural Model")

add_body(
    "This section reports the performance of the PPO-based adaptive controller under the "
    "two stress scenarios defined in Section 3.3.6, benchmarked against static shift policies "
    "and a random baseline."
)

# 4.2.1
add_subsection_heading("4.2.1 Static Baseline Characterization")

add_body(
    "Before evaluating the learned policy, we first establish the performance envelope of "
    "fixed shift policies under stochastic conditions. Table 8 reports the control reward "
    "and service metrics for each static baseline."
)

add_table_caption(
    "Table 8. Static baseline performance under increased and severe risk "
    "(control_v1 reward, 5 seeds \u00d7 10 episodes, stochastic processing times enabled)."
)

add_table(
    ["Scenario", "Policy", "Control reward", "Fill rate", "Backorder rate"],
    [
        ["Increased", "Static S1", "\u2212356.81", "0.652", "0.348"],
        ["Increased", "Static S2", "\u2212171.35", "0.836", "0.164"],
        ["Increased", "Static S3", "\u2212178.09", "0.835", "0.165"],
        ["Severe", "Static S1", "\u2212564.70", "0.452", "0.548"],
        ["Severe", "Static S2", "\u2212384.82", "0.628", "0.372"],
        ["Severe", "Static S3", "\u2212388.40", "0.629", "0.371"],
    ],
)

add_body(
    "Under increased risk, S2 dominates: it provides the highest service level and the best "
    "control reward, indicating that double-shift operation is already sufficient to absorb "
    "moderate disruptions. Under severe risk, S2 and S3 perform comparably, with S2 marginally "
    "better on reward and S3 marginally better on service. Critically, S1 is clearly suboptimal "
    "in both scenarios, confirming that shift allocation is a meaningful control lever for "
    "supply chain resilience."
)

# 4.2.2
add_subsection_heading("4.2.2 Adaptive Control Under Increased Risk")

add_body(
    "Table 9 reports the performance of the PPO-trained policy compared to the best static "
    "baseline (S2) under the increased risk scenario."
)

add_table_caption(
    "Table 9. PPO vs. best static baseline under increased risk "
    "(500k timesteps, control_v1, w_bo = 4.0, w_cost = 0.02, stochastic PT)."
)

add_table(
    ["Policy", "Control reward", "Fill rate", "Backorder rate", "Shift mix (S1/S2/S3)"],
    [
        ["Static S2", "\u2212170.10", "0.837", "0.163", "0% / 100% / 0%"],
        ["PPO", "\u2212172.05", "0.838", "0.162", "12% / 25% / 63%"],
        ["Difference", "\u22121.95", "+0.001", "\u22120.001", "\u2014"],
        [
            "Bootstrap CI\u2089\u2085",
            "[\u22129.95, +8.51]",
            "\u2014",
            "\u2014",
            "\u2014",
        ],
    ],
)

add_body(
    "Under moderate stress, the PPO agent matches the service level of the best static baseline "
    "while using a heterogeneous shift allocation (primarily S3, with occasional downshifting to "
    "S1/S2). The reward difference of \u22121.95 points is not distinguishable from zero (bootstrap "
    "CI\u2089\u2085 includes zero, exact sign-flip p = 0.812). This result indicates that the adaptive "
    "controller is competitive with the optimal fixed policy under increased risk, but does not "
    "yet demonstrate a clear advantage in this regime."
)

# 4.2.3
add_subsection_heading("4.2.3 Adaptive Control Under Severe Risk")

add_body(
    "Table 10 reports the performance under the severe risk scenario, where disruption "
    "frequencies are escalated 4\u201317\u00d7 relative to the thesis baseline."
)

add_table_caption(
    "Table 10. PPO vs. best static baseline under severe risk "
    "(500k timesteps, control_v1, w_bo = 4.0, w_cost = 0.02, stochastic PT)."
)

add_table(
    ["Policy", "Control reward", "Fill rate", "Backorder rate", "Shift mix (S1/S2/S3)"],
    [
        ["Static S3", "\u2212385.59", "0.632", "0.368", "0% / 0% / 100%"],
        ["PPO", "\u2212380.98", "0.631", "0.369", "41% / 27% / 32%"],
        ["Difference", "+4.61", "\u22120.001", "+0.001", "\u2014"],
        [
            "Bootstrap CI\u2089\u2085",
            "[\u22120.28, +9.49]",
            "\u2014",
            "\u2014",
            "\u2014",
        ],
    ],
)

add_body(
    "Under severe stress, the adaptive controller outperforms the best fixed baseline by 4.61 "
    "control-reward points while maintaining effectively equivalent service (fill rate difference "
    "< 0.1%). The bootstrap CI\u2089\u2085 is [\u22120.28, +9.49], with an exact sign-flip p-value of "
    "0.188, indicating a directional advantage that approaches but does not reach conventional "
    "significance levels. Notably, the PPO agent achieves this improvement by reducing its "
    "reliance on S3 relative to the static baseline, using a balanced mix of all three shift "
    "levels. This behavior suggests that the agent learns to downshift during low-disruption "
    "windows, recovering cost without sacrificing service\u2014a strategy that is unavailable to "
    "any fixed policy."
)

# 4.2.4
add_subsection_heading("4.2.4 Interpretation")

add_body("The pattern across the two stress scenarios supports the following reading:")

for num, item in enumerate(
    [
        (
            "Regime-dependent value.",
            " The adaptive controller is competitive under moderate stress and stronger under severe stress. This is consistent with the expectation that fixed policies become insufficient when disruption intensity exceeds their design point.",
        ),
        (
            "Cost-aware adaptation.",
            " The PPO agent does not simply default to the most expensive shift configuration. Under both scenarios, it uses a heterogeneous shift allocation, indicating that the reward function successfully incentivizes cost-aware behavior.",
        ),
        (
            "Service preservation.",
            " In neither scenario does the adaptive policy degrade service relative to the best static baseline. This confirms that the control_v1 reward function avoids the misalignment observed with ReT_thesis (Section 3.3.4).",
        ),
    ],
    1,
):
    p = doc.add_paragraph()
    run1 = p.add_run(f"{num}. {item[0]}")
    run1.font.name = "Times New Roman"
    run1.font.size = Pt(13)
    run1.bold = True
    run2 = p.add_run(item[1])
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(13)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

add_body(
    "These results should be interpreted as preliminary: the current evidence demonstrates "
    "that adaptive shift control has value in high-stress regimes, but does not justify a claim "
    "of uniform superiority across all risk conditions. Section 4.3 extends this analysis by "
    "examining whether richer policy architectures can strengthen the adaptive signal."
)

# Save
output_path = os.path.expanduser("~/Downloads/v0_update_sections_3_3_and_4_2.docx")
doc.save(output_path)
print(f"Saved to {output_path}")

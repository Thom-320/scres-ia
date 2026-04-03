#!/usr/bin/env python3
"""Prepare a minimal team handoff packet for the v0 manuscript.

This script creates two sendable artifacts under ``output/doc``:

1. A lightly patched copy of the legacy v0 manuscript with only the DES-side
   additions that Thom owns.
2. A literature-links document derived from ``docs/for_team/literature_links.md``.

The original v0 DOCX is never modified in place.
"""

from __future__ import annotations

import argparse
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

LITERATURE_REPLACEMENT = (
    "The literature on reinforcement learning (RL) applied to supply chain "
    "management has grown substantially in recent years. Comprehensive reviews "
    "by Yan et al. (2022) and Rolf et al. (2023) classify a wide range of RL "
    "algorithms and applications in logistics and SCM, noting that inventory "
    "control is the dominant application domain. Rolf et al. (2025) review 72 "
    "papers combining discrete-event simulation (DES) with machine learning, "
    "confirming that DES+RL is the most frequent hybrid paradigm. In the "
    "resilience domain, Ding et al. (2026) use multi-agent RL (MAPPO) for "
    "supply chain resilience reconfiguration under disruption, though their "
    "focus is network topology redesign rather than operational control on a "
    "detailed DES. Garrido et al. (2024) proposed integrating artificial "
    "intelligence algorithms into SCRES simulation models, identifying neural "
    "networks and reinforcement learning as suitable candidates for bridging "
    "the gap between static simulation and adaptive learning. The present work "
    "operationalizes this proposal."
)

DES_INTERFACE_PARAGRAPH = (
    "To enable integration with learning algorithms, the DES was wrapped in a "
    "Gymnasium-compatible interface (Towers et al. 2024). At each decision "
    "epoch, set at 168 simulated hours (one operational week), the interface "
    "exposes a state observation vector derived from the simulation's internal "
    "variables and accepts a continuous action vector that modifies "
    "operational parameters before advancing the simulation by one step. The "
    "interface supports multiple observation versions, ranging from 15 to 46 "
    "dimensions, and action contracts from 5 to 7 continuous dimensions, "
    "allowing systematic comparison of different control scopes over the same "
    "DES backbone."
)

DES_CONCLUSION_PARAGRAPH = (
    "In this way, the DES provides a standalone resilience model and, at the "
    "same time, the operational backbone for the hybrid learning architecture "
    "introduced in the next subsection."
)

TABLE_5_INTRO = (
    "To characterize the DES behavior under disruption, two stochastic risk "
    "configurations were evaluated over the 20-year simulation horizon. "
    "Table 5 summarizes the results."
)

TABLE_5_CAPTION = "Table 5. DES performance under stochastic risk configurations"

TABLE_5_ROWS = [
    [
        "Configuration",
        "Avg. annual delivery",
        "Fill rate",
        "Total backorders",
        "Disruptive events",
    ],
    ["Deterministic (Cf0, S=1)", "733,621", "99.3%", "41", "0"],
    ["Current risk", "677,750", "68.3%", "1,825", "8,247"],
    ["Increased risk", "549,250", "45.6%", "3,132", "13,705"],
]

SEND_NOTE = """# Team Handoff Packet

Suggested send order for Garrido and David:

1. `literature_links_for_garrido_david.docx` or its PDF version
2. `v0_neuralNet-scres_DES-minimal.docx` or its PDF version

Suggested message to paste:

Hi Alex and David,

I am sending two items to re-sync us with the current state of the project.

1. A short literature-links file with the papers we were missing in the WhatsApp discussion, especially Ding et al. (2026), the main RL-in-SCM reviews, and the DES+ML review.
2. A minimally updated v0 draft with only the DES-side additions already consolidated on my side. I did not modify the RL-specific sections, so David can still own that part.

The main literature takeaway is that we should avoid a "first RL for supply chain resilience" claim. A safer and stronger positioning is that we operationalize the agenda proposed in Garrido et al. (2024) on cumulative learning in SCRES, using a thesis-grounded DES for the military food supply chain.
"""


def _insert_paragraph_after(
    paragraph: Paragraph, text: str = "", style: str | None = None
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para


def _insert_table_after(paragraph: Paragraph, rows: int, cols: int) -> Table:
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=Inches(6.5))
    paragraph._p.addnext(table._tbl)
    return table


def _set_table_grid_style(table: Table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass


def _normalize_spaces(text: str) -> str:
    return " ".join(text.split())


def _strip_markdown_inline(text: str) -> str:
    return text.replace("**", "").replace("*", "")


def _find_paragraph(doc: Document, startswith: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if _normalize_spaces(paragraph.text).startswith(startswith):
            return paragraph
    raise ValueError(f"Could not find paragraph starting with: {startswith}")


def _has_table_caption(doc: Document, caption: str) -> bool:
    return any(_normalize_spaces(p.text) == caption for p in doc.paragraphs)


def _patch_v0_docx(source_docx: Path, output_docx: Path) -> None:
    doc = Document(source_docx)

    placeholder = _find_paragraph(doc, "XXX et al. (2025)")
    placeholder.text = LITERATURE_REPLACEMENT

    warmup = _find_paragraph(doc, "The simulation horizon spans 20 years")
    if "Gymnasium-compatible interface" not in warmup.text:
        warmup.text = (
            "The simulation horizon spans 20 years (161,280 hours under the "
            "thesis basis of 8,064 hours per year). Following the original "
            "validation logic, the warm-up period ends when the first batch of "
            "5,000 rations reaches Op9, which occurs at approximately 838.8 "
            "hours under deterministic conditions; only after that point are "
            "steady-state observations interpreted."
        )
        inserted = _insert_paragraph_after(
            warmup, DES_INTERFACE_PARAGRAPH, style=warmup.style.name
        )
        _insert_paragraph_after(
            inserted, DES_CONCLUSION_PARAGRAPH, style=warmup.style.name
        )

    if not _has_table_caption(doc, TABLE_5_CAPTION):
        anchor = _find_paragraph(
            doc,
            "Under deterministic Cf0 conditions with risks disabled",
        )
        intro = _insert_paragraph_after(anchor, TABLE_5_INTRO, style=anchor.style.name)
        caption = _insert_paragraph_after(
            intro, TABLE_5_CAPTION, style=anchor.style.name
        )
        table = _insert_table_after(
            caption, rows=len(TABLE_5_ROWS), cols=len(TABLE_5_ROWS[0])
        )
        _set_table_grid_style(table)
        for row_idx, row in enumerate(TABLE_5_ROWS):
            for col_idx, value in enumerate(row):
                cell = table.cell(row_idx, col_idx)
                cell.text = value
                if row_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def _markdown_to_docx(markdown_path: Path, output_docx: Path) -> None:
    doc = Document()
    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped == "---":
            continue
        if stripped.startswith("# "):
            doc.add_heading(_strip_markdown_inline(stripped[2:]), level=1)
            continue
        if stripped.startswith("## "):
            doc.add_heading(_strip_markdown_inline(stripped[3:]), level=2)
            continue
        match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if match:
            doc.add_paragraph(
                _strip_markdown_inline(match.group(2)),
                style="List Number",
            )
            continue
        if stripped.startswith("* ") or stripped.startswith("- "):
            doc.add_paragraph(_strip_markdown_inline(stripped[2:]), style="List Bullet")
            continue
        doc.add_paragraph(_strip_markdown_inline(stripped))

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def _write_send_note(output_dir: Path) -> None:
    (output_dir / "team_handoff_readme.md").write_text(SEND_NOTE, encoding="utf-8")


def _convert_to_pdf(docx_path: Path, output_dir: Path) -> None:
    subprocess.run(
        [
            "soffice",
            f"-env:UserInstallation=file:///tmp/lo_profile_{docx_path.stem}",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ],
        check=True,
    )


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source-docx",
        type=Path,
        default=Path("/Users/thom/Downloads/v.0_neuralNet-scres(1).docx"),
    )
    parser.add_argument(
        "--literature-md",
        type=Path,
        default=Path("docs/for_team/literature_links.md"),
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=Path("output/doc"),
    )
    args = parser.parse_args()

    output_dir = args.output_dir
    output_dir.mkdir(parents=True, exist_ok=True)

    patched_docx = output_dir / "v0_neuralNet-scres_DES-minimal.docx"
    literature_docx = output_dir / "literature_links_for_garrido_david.docx"

    _patch_v0_docx(args.source_docx, patched_docx)
    _markdown_to_docx(args.literature_md, literature_docx)
    _write_send_note(output_dir)

    _convert_to_pdf(patched_docx, output_dir)
    _convert_to_pdf(literature_docx, output_dir)


if __name__ == "__main__":
    main()

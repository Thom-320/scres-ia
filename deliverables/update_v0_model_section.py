#!/usr/bin/env python3
"""Replace v.0's obsolete Section 3.3 with the detailed DES description, in-place style.

Preserves everything else (title, §1-§2, §3.1 hypotheses, the KAN section, §4-§7) and
v.0's own conventions: Normal style throughout, bold-run paragraphs as headings,
one sentence per paragraph.
"""
import copy
from pathlib import Path

from scripts.external_sources import resolve as resolve_external
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Inches

SRC = resolve_external("v.0_neuralNet-scres.docx")
OUT = Path("deliverables/"
           "v0_neuralNet-scres_DES_section_updated.docx")
FIG = Path("deliverables/figures")

# ---- native Word equations (OMML) generated from LaTeX via pandoc ----------
EQ_LATEX = r"""
$$Re_{j} = \begin{cases} \dfrac{AP_{j}}{LT} & \text{if risk active and } AP_{j} > 0 \\[6pt] \dfrac{0.5}{RP_{j}} & \text{if risk active, } AP_{j} = 0,\ RP_{j} > 0 \\[6pt] 0 & \text{if risk active, no recovery} \\[6pt] 1 - \dfrac{B_{t,j} + U_{t,j}}{j} & \text{if no risk active} \end{cases}$$
"""


def build_equation_paragraphs():
    """Return deep-copyable <w:p> elements containing native m:oMath equations."""
    with tempfile.TemporaryDirectory() as td:
        md = Path(td) / "eq.md"
        dx = Path(td) / "eq.docx"
        md.write_text(EQ_LATEX)
        subprocess.run(["pandoc", str(md), "-o", str(dx)], check=True)
        eq_doc = Document(dx)
        out = []
        for p in eq_doc.paragraphs:
            if p._p.xpath(".//*[local-name()='oMath']"):
                out.append(copy.deepcopy(p._p))
        if not out:
            raise SystemExit("pandoc produced no oMath equations")
        return out


shutil.copy(SRC, OUT)
doc = Document(OUT)
paras = doc.paragraphs

start = next(i for i, p in enumerate(paras)
             if p.text.strip().startswith("3.3 Python Hybrid"))
end = next(i for i, p in enumerate(paras)
           if p.text.strip().startswith("Distributional Kolmogorov"))
anchor = paras[end]
print(f"reemplazando parrafos {start}..{end-1} ({end-start} parrafos)")
for p in paras[start:end]:
    p._element.getparent().remove(p._element)


def before(style=None):
    p = doc.add_paragraph()
    anchor._p.addprevious(p._p)
    return p


def head(text, italic=False):
    p = before()
    r = p.add_run(text)
    r.bold = True
    r.italic = italic
    return p


def sent(text):
    p = before()
    p.add_run(text)
    return p


def figure(path, caption, width=6.3):
    p = before()
    p.alignment = 1
    p.add_run().add_picture(str(path), width=Inches(width))
    c = before()
    c.alignment = 1
    r = c.add_run(caption)
    r.italic = True


head("3.3 The discrete-event simulation model")

head("3.3.1 Model genealogy and Baseline 0", italic=True)
for s in (
    "The starting point is the military food supply-chain (MFSC) simulation reported by Garrido-Rios (2017).",
    "The original model represents procurement, production, transport, storage, and theatre delivery as thirteen linked operations.",
    "The original model evaluates fixed inventory-buffer and manufacturing-capacity scenarios under a long-run stationary design.",
    "It does not make weekly decisions from the current system state.",
    "This original configuration is Baseline 0 in our comparison ladder.",
    "Baseline 0 uses discretionary ex-ante policies and is therefore a conceptual anchor rather than a matched adaptive comparator.",
    "We reconstruct the specification as a Python discrete-event simulation.",
    "The reconstruction preserves the operation topology, the decision tables, the risk parameter tables, and the operational ReT formula.",
    "It also repairs causal links that were absent from early code versions, such as conservation-respecting procurement.",
    "The reconstruction is a high-fidelity transcription of the static specification; complete endogenous replication of every thesis experiment is not claimed.",
): sent(s)

head("3.3.2 The thirteen-operation system", italic=True)
for s in (
    "The simulation advances by scheduled events rather than by a fixed time step.",
    "Each operation can release material, receive material, queue, process, and trigger transport.",
    "Production follows an assemble-to-stock logic from contracting (Op1) to the supply battalion (Op9).",
    "Distribution follows an assemble-to-order logic from the supply battalion to the theatre (Op13).",
    "The promised downstream order lead time is forty-eight hours.",
    "Figure M1 summarizes the thirteen operations with their processing times, order quantities, reorder cadences, and the Garrido-native risks attached to each operation.",
): sent(s)
figure(FIG / "fig1_mfsc_flow.png",
       "Figure M1. Thirteen-operation flow of the MFSC (after Garrido-Rios, Figure 6.2). "
       "Red chips mark the risks affecting each operation.")

head("3.3.3 Material flow and conservation", italic=True)
for s in (
    "A finished lot cannot be created unless the required upstream material and production rights exist.",
    "Each finished lot contains 5,000 ration units.",
    "The two-product extension assigns a product label (P_C or P_H) before the lot enters the controlled production sequence.",
    "The label cannot be overwritten after a twenty-four-hour activation delay, which prevents retroactive renaming of inventory.",
    "Op9 inventory is stored separately per product, and cross-product substitution is disabled.",
    "The order queue is work-conserving and product-feasible: an unavailable product cannot block a feasible order for the other product.",
    "Late orders remain in the ledger as backorders, and the pending-order list is capped at sixty orders, following the source model.",
    "Every experimental arm receives the same scheduled production and transport entitlement.",
    "Mass and product-partition checks verify that units are neither created nor destroyed by the controller.",
): sent(s)

head("3.3.4 Time, warm-up, and episode boundaries", italic=True)
for s in (
    "The thesis reports a deterministic warm-up of 838.8 hours; the reconstruction triggers warm-up on physical readiness instead, because stochastic events can shift the calendar.",
    "The two-product experiment uses a stricter product-balanced warm-up: one real 5,000-unit lot of each product must traverse Op1 through Op9 and coexist at Op9.",
    "No finished inventory is ever injected directly.",
    "The treatment episode lasts eight decision weeks of 168 hours each.",
    "Six demand orders arrive per week at offsets of 30, 54, 78, 102, 126, and 150 hours.",
    "No production or demand rights are created after the eighth decision.",
    "A common 1,344-hour clearance tail lets material and open orders propagate, and every policy is scored at the same final time.",
    "Figure M2 shows the episode structure.",
): sent(s)
figure(FIG / "fig6_timeline.png",
       "Figure M2. Episode structure: product-balanced warm-up, eight weekly decisions "
       "with three 5,000-unit batch completions each, clearance tail, and common score time.")

head("3.3.5 Demand process and partial observability", italic=True)
for s in (
    "Each episode contains forty-eight demand orders with integer quantities drawn from 2,400 to 2,600 rations, exactly the thesis demand range.",
    "The requested product follows a two-state latent Markov regime: one regime makes P_C dominant, the other makes P_H dominant.",
    "The regime persists between demand events with probability rho, and the dominant product is requested with probability s.",
    "The evaluated cells are (rho=0.75, s=0.90), (rho=0.90, s=0.75), and (rho=0.90, s=0.90).",
    "The controller never observes the latent regime and never receives the true rho or s of the evaluation cell.",
    "A fixed hidden-Markov belief filter supplies a deployable estimate with the same fixed parameters in every cell, so no policy receives privileged knowledge of the sensitivity condition.",
    "Demand tapes are generated independently of policy actions, which enables exact paired comparisons.",
    "Future demand, random seeds, tape identifiers, oracle calendars, and terminal outcomes are forbidden observations.",
): sent(s)

head("3.3.6 Decision and observation contracts", italic=True)
for s in (
    "The controller does not choose production volume, shift count, vehicle count, or demand quantity.",
    "It chooses only how the next three fixed 5,000-unit batch targets are divided between P_C and P_H, one action per week from the set {0, 1, 2, 3} equal to the number of P_C batches.",
    "Batch positions are centered (for example, action 1 is P_H, P_C, P_H) to remove intra-week timing advantages.",
    "Actions activate after twenty-four hours, with batch completions at 24, 72, and 120 hours within the week.",
    "Every eight-week calendar therefore commits exactly twenty-four batch slots and 120,000 units of controlled production rights.",
    "The policy observes a normalized twenty-one-dimensional vector at each decision: per-product inventory, locked pipeline, backlog quantity and count, maximum backlog age, in-flight quantity, the fixed HMM belief and predicted share, the previous action (one-hot), and the episode phase.",
    "The learning context is therefore a partially observed weekly allocation problem: the network must infer the demand regime from operational history.",
    "Figure M3 summarizes the experimental architecture and the three controllers evaluated under identical resources.",
): sent(s)
figure(FIG / "fig2_framework.png",
       "Figure M3. Experimental framework: the DES backbone and two-product extension feed a "
       "21-dimensional observation to three controllers -- the exact 65,536-calendar open-loop "
       "frontier, the strongest tested structured feedback family, and RecurrentPPO -- all scored "
       "by the same terminal ReT and service ledger.")

head("3.3.7 Resilience metric and reward", italic=True)
for s in (
    "The primary outcome is the operational ReT formula from Garrido's raw Excel workbooks, reproduced without clipping, normalization, or replacement.",
    "Each order j is scored at its request-time snapshot according to the following piecewise definition:",
): sent(s)
for eq_p in build_equation_paragraphs():
    anchor._p.addprevious(eq_p)
for s in (
    "Here AP is the autotomy period, LT the promised forty-eight-hour lead time, RP the recovery period, and B and U the backorders and unattended orders accumulated at the request time.",
    "The formula was recomputed on 47,546 source-workbook rows with zero mismatches.",
    "The two-product primary experiment runs risk-off, so the no-risk branch is operative, and risk-adaptation claims are explicitly out of scope.",
    "The reward is zero at intermediate decisions and equals the official terminal ReT at the common score time, avoiding any unvalidated shaping of the construct.",
    "Because the visible ReT population contains completed non-lost orders, unresolved demand, lost demand, per-product fill, terminal inventory, and resource use are mandatory companion outcomes.",
    "Figure M4 shows the scoring branches.",
): sent(s)
figure(FIG / "fig5_ret_tree.png",
       "Figure M4. Order-level ReT scoring branches, reproduced exactly from the source "
       "workbooks (47,546 rows, zero mismatches). Program Q operates on the no-risk branch.")

head("3.3.8 Sensitive assumptions", italic=True)
for s in (
    "The most sensitive modeling decisions are those that determine what information is available, when actions become physical, and which orders enter the metric.",
    "Equal-time events follow a strictly half-open rule: a lot arriving exactly at a release timestamp is not eligible for that release, and an order completion at a request time is processed before the request snapshot.",
    "These conventions prevent the policy and the metric from seeing zero-time information that would not be operationally available.",
    "The product classes are synthetic and share identical physics, so the experiment isolates allocation under demand uncertainty rather than product-specific manufacturing complexity.",
    "Storage is unlimited, vehicle capacity is a fixed entitlement, and substitution is disabled.",
    "A finite horizon could invite end-of-episode inventory hoarding; the fixed production entitlement, the clearance tail, and the terminal product ledger expose and bound this behavior, and no reward is assigned to ending inventory.",
): sent(s)

doc.save(OUT)
print("guardado:", OUT)
d2 = Document(OUT)
print("parrafos ahora:", len(d2.paragraphs), "| imagenes:",
      sum(1 for _ in d2.inline_shapes))

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    "deliverables/"
    "Program_Q_DES_Model_Description_CIE_reviewed.docx"
)
FIG1 = Path(
    "deliverables/"
    "fig1_flow_reviewed.png"
)

BLUE = "2E74B5"
DARK_BLUE = "17365D"
LIGHT_BLUE = "EEF5FA"
LIGHT_GRAY = "F4F6F9"
MID_GRAY = "D9E1E8"
DARK_GRAY = "4F5B66"
WHITE = "FFFFFF"
BLACK = "1A1A1A"
AMBER = "FFF2CC"
GREEN = "E2F0D9"
RED = "FCE4D6"

PAGE_WIDTH_DXA = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = PAGE_WIDTH_DXA) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.line_spacing = line


def set_font(run, name="Calibri", size=11, bold=False, color=BLACK, italic=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_border_bottom(paragraph, color=MID_GRAY, size=6, space=1) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


def add_sentence(doc: Document, text: str, style: str = "Body Text", *, keep=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = keep
    paragraph.paragraph_format.keep_together = True
    return paragraph


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Caption")
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_equation(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Equation")
    run = paragraph.add_run(text)
    set_font(run, name="Cambria Math", size=11, color=DARK_BLUE)
    return paragraph


def add_callout(doc: Document, title: str, sentences: Sequence[str], fill=LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=100, end=160)
    p = cell.paragraphs[0]
    p.style = doc.styles["Callout Title"]
    p.add_run(title)
    for sentence in sentences:
        p = cell.add_paragraph(style="Callout Text")
        p.add_run(sentence)
    doc.add_paragraph()


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    widths: Sequence[int],
    *,
    compact=False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for cell, value, width in zip(header.cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, DARK_BLUE)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        run = paragraph.add_run(value)
        set_font(run, size=9 if compact else 9.5, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for cell, value, width in zip(row.cells, values, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if row_index % 2 == 1:
                set_cell_shading(cell, LIGHT_GRAY)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            set_paragraph_spacing(paragraph, after=0, line=1.05)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(value)
            set_font(run, size=8.4 if compact else 9.0, color=BLACK)
    doc.add_paragraph()


def add_status_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table)
    headers = ["Status", "Statement", "Required wording"]
    widths = [1200, 3900, 4260]
    for cell, text, width in zip(table.rows[0].cells, headers, widths):
        set_cell_width(cell, width)
        set_cell_shading(cell, DARK_BLUE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.0)
        r = p.add_run(text)
        set_font(r, size=9.2, bold=True, color=WHITE)
    for status, statement, wording in rows:
        row = table.add_row()
        prevent_row_split(row)
        colors = {"SUPPORTED": GREEN, "CONDITIONAL": AMBER, "NOT SUPPORTED": RED}
        for index, (cell, text, width) in enumerate(
            zip(row.cells, (status, statement, wording), widths)
        ):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            if index == 0:
                set_cell_shading(cell, colors.get(status, LIGHT_GRAY))
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.05)
            r = p.add_run(text)
            set_font(r, size=8.8, bold=(index == 0))
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    fmt = normal.paragraph_format
    fmt.space_after = Pt(8)
    fmt.line_spacing = 1.333
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT

    body = styles["Body Text"]
    body.font.name = "Calibri"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    body.font.size = Pt(11)
    body.font.color.rgb = RGBColor.from_string(BLACK)
    body.paragraph_format.space_after = Pt(8)
    body.paragraph_format.line_spacing = 1.333
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body.paragraph_format.widow_control = True
    body.paragraph_format.keep_together = True

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.bold = True
    caption.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True

    for style_name in ("Callout Title", "Callout Text", "Equation", "Small Note"):
        if style_name not in styles:
            styles.add_style(style_name, 1)

    callout_title = styles["Callout Title"]
    callout_title.font.name = "Calibri"
    callout_title._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    callout_title.font.size = Pt(11)
    callout_title.font.bold = True
    callout_title.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout_title.paragraph_format.space_after = Pt(4)

    callout_text = styles["Callout Text"]
    callout_text.font.name = "Calibri"
    callout_text._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    callout_text.font.size = Pt(9.5)
    callout_text.font.color.rgb = RGBColor.from_string(BLACK)
    callout_text.paragraph_format.space_after = Pt(4)
    callout_text.paragraph_format.line_spacing = 1.15

    equation = styles["Equation"]
    equation.font.name = "Cambria Math"
    equation._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    equation.font.size = Pt(11)
    equation.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    equation.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    equation.paragraph_format.space_before = Pt(4)
    equation.paragraph_format.space_after = Pt(8)
    equation.paragraph_format.keep_together = True

    small_note = styles["Small Note"]
    small_note.font.name = "Calibri"
    small_note._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    small_note.font.size = Pt(8.5)
    small_note.font.italic = True
    small_note.font.color.rgb = RGBColor.from_string(DARK_GRAY)
    small_note.paragraph_format.space_after = Pt(4)
    small_note.paragraph_format.line_spacing = 1.05


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, after=2, line=1.0)
    r = p.add_run("SCRES-IA  |  DES model and learning environment")
    set_font(r, size=8.5, bold=True, color=DARK_GRAY)
    set_paragraph_border_bottom(p, color=MID_GRAY, size=5)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(p, after=0, line=1.0)
    r = p.add_run("Working insert for C&IE  |  ")
    set_font(r, size=8, color=DARK_GRAY)
    add_field(p, "PAGE")


def build_reviewed_flow_figure() -> None:
    """Build a clean, publication-style topology without overlapping labels."""

    fig, ax = plt.subplots(figsize=(13.0, 6.1))
    ax.set_xlim(-0.8, 12.5)
    ax.set_ylim(-0.1, 6.7)
    ax.axis("off")

    dark = "#17365D"
    blue = "#2E74B5"
    light = "#EEF5FA"
    transport = "#DCEAF5"
    production = "#E2F0D9"
    demand = "#FFF2CC"
    muted = "#4F5B66"
    width = 2.55
    height = 1.0

    positions = {
        1: (0.0, 5.35),
        2: (3.05, 5.35),
        3: (6.10, 5.35),
        4: (9.15, 5.35),
        5: (9.15, 3.55),
        6: (6.10, 3.55),
        7: (3.05, 3.55),
        8: (0.0, 3.55),
        9: (0.0, 1.75),
        10: (3.05, 1.75),
        11: (6.10, 1.75),
        12: (9.15, 1.75),
        13: (9.15, 0.05),
    }
    labels = {
        1: ("Op1 — MLA contracting", "PT 672 h · ROP 4,032 h"),
        2: ("Op2 — Suppliers ship raw material", "PT 24 h · monthly kit"),
        3: ("Op3 — WDC receive/store", "PT 24 h · weekly 15,500/rm"),
        4: ("Op4 — LOC to assembly", "PT 24 h"),
        5: ("Op5 — Pre-assembly", "320.5 rations/h"),
        6: ("Op6 — Assembly", "balanced line"),
        7: ("Op7 — QC and packaging", "5,000-unit batch · 48 h"),
        8: ("Op8 — LOC to supply battalion", "PT 24 h"),
        9: ("Op9 — Supply battalion store", "daily 2,400–2,600"),
        10: ("Op10 — LOC to CSSU", "PT 24 h"),
        11: ("Op11 — CSSU issue", "PT 0 h · 2 CSSUs"),
        12: ("Op12 — LOC to theatre", "PT 24 h"),
        13: ("Op13 — Theatre demand", "U(2,400–2,600)/day × 6"),
    }
    production_ops = {5, 6, 7}
    transport_ops = {4, 8, 10, 12}

    for operation, (x, y) in positions.items():
        if operation == 13:
            fill = demand
        elif operation in production_ops:
            fill = production
        elif operation in transport_ops:
            fill = transport
        else:
            fill = light
        ax.add_patch(
            FancyBboxPatch(
                (x, y),
                width,
                height,
                boxstyle="round,pad=0.018",
                facecolor=fill,
                edgecolor=dark,
                linewidth=1.4,
            )
        )
        title, detail = labels[operation]
        ax.text(
            x + width / 2,
            y + 0.70,
            title,
            ha="center",
            va="center",
            fontsize=9.0,
            fontweight="bold",
            color=dark,
        )
        ax.text(
            x + width / 2,
            y + 0.30,
            detail,
            ha="center",
            va="center",
            fontsize=8.0,
            color="#1A1A1A",
        )

    chain = list(range(1, 14))
    for source, target in zip(chain[:-1], chain[1:]):
        x1, y1 = positions[source]
        x2, y2 = positions[target]
        if y1 == y2:
            if x2 > x1:
                start = (x1 + width, y1 + height / 2)
                end = (x2, y2 + height / 2)
            else:
                start = (x1, y1 + height / 2)
                end = (x2 + width, y2 + height / 2)
        else:
            start = (x1 + width / 2, y1)
            end = (x2 + width / 2, y2 + height)
        ax.add_patch(
            FancyArrowPatch(
                start,
                end,
                arrowstyle="-|>",
                mutation_scale=14,
                linewidth=1.4,
                color=blue,
                shrinkA=0,
                shrinkB=0,
            )
        )

    ax.plot([-0.35, -0.35], [1.70, 6.38], color="#2F9E77", linewidth=3)
    ax.text(
        -0.53,
        4.04,
        "Assemble-to-stock · Op1–Op9",
        rotation=90,
        ha="center",
        va="center",
        fontsize=8.0,
        fontweight="bold",
        color="#2F9E77",
    )
    ax.plot([-0.35, -0.35], [0.0, 2.70], color=blue, linewidth=3)
    ax.text(
        -0.53,
        1.35,
        "Assemble-to-order · Op9–Op13",
        rotation=90,
        ha="center",
        va="center",
        fontsize=8.0,
        fontweight="bold",
        color=blue,
    )
    ax.text(
        0,
        -0.03,
        "The Program Q primary experiment uses this topology with active risks disabled.",
        ha="left",
        va="top",
        fontsize=7.5,
        color=muted,
        style="italic",
    )
    fig.savefig(FIG1, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def build_document() -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    props = doc.core_properties
    props.title = "Discrete-Event Simulation Model and Learning Environment"
    props.subject = "Program Q manuscript-ready replacement section for Computers & Industrial Engineering"
    props.author = "Thomas Chisica and SCRES-IA collaborators"
    props.keywords = "supply chain resilience, discrete-event simulation, RecurrentPPO, structured feedback, ReT"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_before = Pt(88)
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Discrete-Event Simulation Model")
    set_font(run, size=23, bold=True, color=DARK_BLUE)
    run.add_break()
    run2 = title.add_run("and Learning Environment")
    set_font(run2, size=23, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run(
        "Manuscript-ready replacement for the model and learning-context sections"
    )
    set_font(run, size=13, color=BLUE)

    add_callout(
        doc,
        "Purpose of this insert",
        [
            "This document describes the DES sequentially, with one sentence per paragraph, as requested during the 22 July 2026 meeting.",
            "The text is written for direct integration into a Computers & Industrial Engineering manuscript.",
            "The document separates the Garrido reference model, the Program Q extension, the controller contract, and the current evidence boundary.",
        ],
        fill=LIGHT_BLUE,
    )

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(meta, 7200)
    meta_values = [
        ("Target journal", "Computers & Industrial Engineering"),
        ("Scientific lane", "Submission A — Program Q"),
        ("Prepared for", "Coauthor review and shared-document integration"),
        ("Version date", date(2026, 7, 26).strftime("%d %B %Y")),
    ]
    for row, (label, value) in zip(meta.rows, meta_values):
        set_cell_width(row.cells[0], 2100)
        set_cell_width(row.cells[1], 5100)
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell in row.cells:
            set_cell_margins(cell)
        p = row.cells[0].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.0)
        r = p.add_run(label)
        set_font(r, size=9, bold=True, color=DARK_BLUE)
        p = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p, after=0, line=1.0)
        r = p.add_run(value)
        set_font(r, size=9)

    doc.add_page_break()

    add_heading(doc, "Editorial control note", 1)
    for sentence in (
        "The two supplied v0 DOCX files are byte-identical copies of the same unfinished manuscript.",
        "The v0 manuscript contains placeholders, an obsolete five-dimensional action space, an obsolete fifteen-dimensional observation, and claims that precede the executed Program Q evidence.",
        "This insert should replace the v0 model and learning-environment prose instead of being appended to it without reconciliation.",
        "The final manuscript should keep Program Q and the later retained-learning Q-R1 study as separate scientific products.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "Recommended paper position", 1)
    for sentence in (
        "The publication question is whether recurrent feedback adds value beyond an exhaustive static frontier and whether that value requires a neural controller.",
        "Program Q answers the first part positively and the second part negatively within its tested contract.",
        "The defensible result is therefore feedback value with structured-control equivalence rather than neural superiority.",
        "The model description must make this decomposition visible before the results are presented.",
    ):
        add_sentence(doc, sentence)

    add_callout(
        doc,
        "Central claim boundary",
        [
            "Program Q outperforms every open-loop calendar in its disclosed two-product extension.",
            "Program Q is practically equivalent to the strongest tested structured feedback family.",
            "Program Q does not establish a neural premium, worst-product safety, cumulative learning, or improvement under active Garrido-native risks.",
        ],
        fill=AMBER,
    )

    add_heading(doc, "1. Model genealogy and scope", 1)
    for sentence in (
        "The starting point is the military food supply-chain simulation reported by Garrido-Rios.",
        "The original model represents the procurement, production, transport, storage, and theatre-delivery system as thirteen linked operations.",
        "The original model evaluates fixed inventory-buffer and manufacturing-capacity scenarios under a long-run stationary design.",
        "The original model does not make weekly decisions from the current system state.",
        "The original model therefore provides the physical reference and the initial static-policy baseline.",
        "The repository reconstructs this specification as a Python discrete-event simulation.",
        "The reconstruction preserves the operation topology, the decision tables, the risk parameter tables, and the operational ReT formula.",
        "The reconstruction also repairs causal links that were absent from early code versions.",
        "The reconstruction is a high-fidelity transcription of the static specification, but it is not yet a validated endogenous reproduction of every thesis experiment.",
        "Program Q adds a disclosed two-product decision problem to this reconstructed backbone.",
        "Program Q disables active risks so that the causal value of product-mix feedback can be identified before risk adaptation is claimed.",
        "The two product classes are synthetic nonfungible mission-suitability classes named P_C and P_H.",
        "The two classes share the same bill of materials, production rate, mass, processing time, and transport entitlement.",
        "The only controlled difference is the product identity required by demand.",
        "This design creates a constrained allocation problem without allowing the controller to buy additional resources.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 1. What is inherited, reconstructed, and added.")
    add_table(
        doc,
        ["Layer", "What the layer contributes", "Scientific limit"],
        [
            (
                "Garrido reference",
                "Thirteen-operation topology, fixed scenario logic, risk definitions, demand quantities, capacities, lead-time semantics, and ReT construct.",
                "The original model is a one-product static-scenario study and is not a weekly adaptive decision benchmark.",
            ),
            (
                "Python reconstruction",
                "Executable event logic, explicit order and resource ledgers, causal procurement and transport repairs, mass checks, and formula replay.",
                "Held-out queue and recovery-tail validation remains incomplete, so full endogenous thesis replication is not claimed.",
            ),
            (
                "Program Q extension",
                "Two nonfungible product classes, eight weekly product-mix decisions, exact open-loop frontier, structured feedback, and recurrent RL.",
                "The product mix is researcher-defined, risks are off, and the result does not represent the full military supply chain.",
            ),
        ],
        [1680, 4100, 3580],
    )

    add_heading(doc, "2. Baseline 0 and the fair comparison ladder", 1)
    for sentence in (
        "Baseline 0 is the original Garrido simulation with fixed ex-ante policies.",
        "Baseline 0 is essential because it anchors the topology, operational assumptions, and resilience metric.",
        "Baseline 0 is not a fair one-to-one performance comparator for Program Q because it uses one homogeneous product and a different decision space.",
        "The fair static comparator inside Program Q is the complete set of 65,536 eight-week product-mix calendars.",
        "Each calendar fixes all eight decisions before the episode begins.",
        "The static frontier therefore contains every open-loop policy available under the Program Q action contract.",
        "The next comparison layer is the strongest tested structured feedback family.",
        "The structured family sees deployable state information and can change its action during the episode.",
        "The final executed comparison layer is RecurrentPPO with a recurrent policy and a multilayer perceptron.",
        "DMLPA and KAN are architecture sidecars until they are rerun under the identical Program Q contract.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 2. Comparator ladder and interpretation.")
    add_table(
        doc,
        ["Level", "Comparator", "Role", "Permitted conclusion"],
        [
            (
                "0",
                "Garrido static reference",
                "Physical and conceptual anchor.",
                "Shows the original static decision logic, but not a matched adaptive comparison.",
            ),
            (
                "1",
                "Exact 65,536-calendar frontier",
                "Primary open-loop benchmark under identical Program Q resources and physics.",
                "A win identifies the value of within-episode feedback.",
            ),
            (
                "2",
                "Structured feedback family",
                "Non-neural adaptive benchmark with the same deployable history.",
                "A comparison separates feedback value from a neural premium.",
            ),
            (
                "3",
                "RecurrentPPO with MLP-LSTM policy",
                "Executed learned feedback controller.",
                "A win over Level 1 establishes learned adaptation, while equivalence with Level 2 does not establish neural superiority.",
            ),
            (
                "4",
                "DMLPA-PPO and KAN-PPO",
                "Architecture ablations under a future matched Program Q run.",
                "No Program Q performance claim is permitted until environment, tapes, reward, budget, and evaluation are matched.",
            ),
        ],
        [700, 2100, 2900, 3660],
        compact=True,
    )

    add_heading(doc, "3. The thirteen-operation discrete-event system", 1)
    for sentence in (
        "The simulation advances by scheduled events rather than by a fixed numerical time step.",
        "Each operation can release material, receive material, create a queue, begin processing, complete processing, or trigger transport.",
        "The event calendar determines the order in which these state changes occur.",
        "The model records material and order identities throughout the flow.",
        "The production segment follows an assemble-to-stock logic through the supply battalion.",
        "The distribution segment follows an assemble-to-order logic from the supply battalion to the theatre.",
        "The promised downstream order lead time is forty-eight hours after the warm-up condition is met.",
    ):
        add_sentence(doc, sentence)

    if FIG1.exists():
        add_caption(doc, "Figure 1. Thirteen-operation flow of the MFSC (after thesis Figure 6.2).")
        pic_paragraph = doc.add_paragraph()
        pic_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_run = pic_paragraph.add_run()
        pic_run.add_picture(str(FIG1), width=Inches(6.5))

    add_caption(doc, "Table 3. Sequential operation map.")
    operations = [
        ("Op1", "Contract suppliers", "12 supplier contracts.", "672 h processing and 4,032 h renewal."),
        ("Op2", "Prepare and ship raw material", "Twelve raw-material streams.", "24 h transport and 190,000 units per stream every 672 h."),
        ("Op3", "Receive and store at WDC", "Raw-material inventory.", "24 h receipt and 15,500 units per stream released every 168 h."),
        ("Op4", "Transport to assembly", "Weekly raw-material kit.", "24 h route time."),
        ("Op5", "Pre-assemble", "Ration units.", "Approximately 320.5 rations per hour and 2,564 units per eight-hour shift."),
        ("Op6", "Assemble", "Ration units.", "Balanced at the Op5 line rate."),
        ("Op7", "Inspect and package", "Finished ration lots.", "Boxes of ten and 5,000-unit release batches every 48 h."),
        ("Op8", "Transport to supply battalion", "Finished 5,000-unit lots.", "24 h transport."),
        ("Op9", "Receive and store", "Finished-ration inventory by product.", "24 h receipt and daily downstream release opportunities."),
        ("Op10", "Transport to CSSU", "Orders of 2,400–2,600 rations.", "24 h transport."),
        ("Op11", "Receive and issue at CSSU", "Two CSSU destinations in the source model.", "Handling is modeled as zero hours because the source reports less than one hour."),
        ("Op12", "Transport to theatre", "Fulfilled order payload.", "24 h transport."),
        ("Op13", "Generate theatre demand", "Regular and contingent orders.", "Regular demand is 2,400–2,600 rations per day on six days each week."),
    ]
    add_table(
        doc,
        ["Operation", "Function", "Flow object", "Source cadence or capacity"],
        operations,
        [700, 2450, 2450, 3760],
        compact=True,
    )

    add_heading(doc, "4. Material flow, queues, and conservation", 1)
    for sentence in (
        "The upstream flow begins with twelve raw-material streams.",
        "The reconstruction can represent these streams as a kit-equivalent replenishment process when the product-mix experiment is active.",
        "A finished lot cannot be created unless the required upstream material and production rights exist.",
        "Each finished lot contains 5,000 ration units.",
        "Program Q assigns a product label before the lot enters the controlled production sequence.",
        "The assigned label cannot be overwritten after its twenty-four-hour activation delay.",
        "The no-relabeling rule prevents the controller from using future demand to rename inventory retroactively.",
        "The Op9 inventory is stored separately for P_C and P_H.",
        "An order can be released only when the requested product is physically available.",
        "Cross-product substitution is disabled.",
        "The queue is work-conserving and product-feasible.",
        "An unavailable product at the head of the policy-independent priority order cannot block a feasible order for the other product.",
        "Late orders remain in the ledger as backorders.",
        "The pending-order list is capped at sixty orders, following the source model.",
        "Every experimental arm receives the same scheduled production and transport entitlement.",
        "Mass and product partition checks verify that units are neither created nor destroyed by the controller.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "5. Time, warm-up, and episode boundaries", 1)
    for sentence in (
        "The original thesis reports a deterministic warm-up of 838.8 hours.",
        "The Python reconstruction can instead trigger warm-up when the first completed lot reaches Op9 because stochastic disruptions can shift the calendar.",
        "Program Q uses a stricter product-balanced warm-up condition.",
        "One real 5,000-unit P_C lot and one real 5,000-unit P_H lot must traverse Op1 through Op9 and coexist at Op9.",
        "No finished inventory is injected directly into the system.",
        "The treatment episode then lasts eight decision weeks.",
        "A decision is made every 168 hours.",
        "Six demand orders arrive during each decision week.",
        "The order offsets are 30, 54, 78, 102, 126, and 150 hours after the weekly decision boundary.",
        "No new production or demand rights are created after the eighth decision.",
        "A common 1,344-hour clearance tail allows material and open orders to propagate after the treatment window.",
        "All policies are scored at the same final time.",
        "Unresolved orders remain visible in the companion ledger at that score time.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "6. Demand process and partial observability", 1)
    for sentence in (
        "Each Program Q episode contains forty-eight demand orders.",
        "Each order quantity is sampled as an integer from 2,400 to 2,600 rations.",
        "The product label is generated by a two-state latent Markov regime.",
        "One regime makes P_C dominant and the other regime makes P_H dominant.",
        "The regime remains unchanged from one demand event to the next with probability ρ.",
        "The dominant product is requested with probability s.",
        "The three evaluated cells are (ρ=0.75, s=0.90), (ρ=0.90, s=0.75), and (ρ=0.90, s=0.90).",
        "The controller never observes the latent regime.",
        "The controller also never receives the true ρ or s of the evaluation cell.",
        "A fixed hidden-Markov belief model supplies a deployable estimate using the same fixed parameters in every cell.",
        "This fixed-model rule prevents the policy from receiving privileged knowledge of the sensitivity condition.",
        "Product-label tapes are generated independently of policy actions.",
        "Future demand, random seeds, tape identifiers, oracle calendars, and terminal outcomes are forbidden observations.",
    ):
        add_sentence(doc, sentence)

    add_callout(
        doc,
        "Why recurrence is justified",
        [
            "The current observation does not reveal the latent demand regime.",
            "Several distinct histories can produce similar inventory and backlog snapshots.",
            "A recurrent controller can use the sequence of observations and actions to summarize evidence that is absent from a single snapshot.",
        ],
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "7. Program Q decision contract", 1)
    for sentence in (
        "The controller does not choose production volume, shift count, vehicle count, or demand quantity.",
        "The controller chooses only how the next three fixed 5,000-unit batch targets are divided between P_C and P_H.",
        "Action zero assigns all three targets to P_H.",
        "Action one assigns the centered sequence P_H, P_C, and P_H.",
        "Action two assigns the centered sequence P_C, P_H, and P_C.",
        "Action three assigns all three targets to P_C.",
        "The centered schedule reduces a timing advantage that would otherwise arise from always placing the minority product first or last.",
        "Each action activates after twenty-four hours.",
        "The associated batch completion offsets are twenty-four, seventy-two, and one hundred twenty hours within the week.",
        "Every eight-week calendar commits exactly twenty-four controlled batch slots.",
        "Every calendar therefore receives exactly 120,000 units of controlled production rights.",
        "The action changes product allocation but cannot expand the resource envelope.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 4. Action mapping within one week.")
    add_table(
        doc,
        ["Action", "Batch 1", "Batch 2", "Batch 3", "Interpretation"],
        [
            ("0", "P_H", "P_H", "P_H", "Allocate all three targets to P_H."),
            ("1", "P_H", "P_C", "P_H", "Allocate one centered target to P_C."),
            ("2", "P_C", "P_H", "P_C", "Allocate two centered targets to P_C."),
            ("3", "P_C", "P_C", "P_C", "Allocate all three targets to P_C."),
        ],
        [800, 1250, 1250, 1250, 4810],
    )

    add_heading(doc, "8. Same-time event convention", 1)
    for sentence in (
        "Discrete-event models require an explicit rule when multiple events share a timestamp.",
        "Program Q uses a strictly half-open information rule.",
        "A lot that reaches Op9 exactly at a daily release timestamp is not eligible for that release.",
        "The lot becomes eligible at the next downstream slot.",
        "An order completion or loss occurring exactly at a new request time is processed before the new request snapshot.",
        "The current request is excluded from its own backlog snapshot.",
        "These rules prevent the policy and the metric from seeing zero-time information that would not have been operationally available.",
        "Scheduler-invariance tests check that results do not depend on an accidental software ordering of equal-time events.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "9. Observation supplied to the controller", 1)
    for sentence in (
        "The policy receives a normalized twenty-one-dimensional observation at each weekly decision.",
        "All components are available from the simulated operational history at the decision time.",
        "Continuous components are scaled to comparable ranges and clipped to the interval from zero to one.",
        "The previous action is represented as a five-position one-hot vector that includes the initial no-action state.",
        "The observation contains no future event or latent-state field.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 5. Twenty-one-dimensional Program Q observation.")
    obs_rows = [
        ("On-hand inventory", "P_C and P_H", "2", "Divide each quantity by 120,000."),
        ("Locked production pipeline", "P_C and P_H", "2", "Divide each quantity by 120,000."),
        ("Backlog quantity", "P_C and P_H", "2", "Divide each quantity by 120,000."),
        ("Backlog order count", "P_C and P_H", "2", "Divide each count by 48."),
        ("Maximum backlog age", "P_C and P_H", "2", "Divide each age by 1,344 hours."),
        ("In-flight quantity", "P_C and P_H", "2", "Divide each quantity by 120,000."),
        ("HMM belief", "Probability of P_C-dominant regime", "1", "Use the fixed deployable filter."),
        ("Predicted demand share", "Predicted P_C share", "1", "Use the fixed deployable filter."),
        ("Previous action", "None or actions 0–3", "5", "Use one-hot encoding."),
        ("Episode phase", "Current week and decisions remaining", "2", "Scale by seven and eight respectively."),
    ]
    add_table(
        doc,
        ["Field group", "Product or meaning", "Dimensions", "Normalization or rule"],
        obs_rows,
        [2300, 2900, 1100, 3060],
        compact=True,
    )

    add_heading(doc, "10. Resilience metric and reward", 1)
    for sentence in (
        "The primary outcome is the operational ReT formula used in Garrido's raw Excel workbooks.",
        "The formula is reproduced without clipping, normalization, or replacement.",
        "When a risk is active and the autotomy period is positive, the order score is the autotomy period divided by promised lead time.",
        "When a risk is active without positive autotomy but with positive recovery, the order score is 0.5 divided by the recovery period.",
        "When a risk is active without recovery, the order score is zero.",
        "When no risk is active, the order score is one minus the accumulated backorders and unattended orders divided by the order index.",
        "Program Q disables active risks, so the no-risk branch is the operative branch in the primary experiment.",
    ):
        add_sentence(doc, sentence)

    add_equation(doc, "Re_j = AP_j / LT_j,  when risk is active and AP_j > 0.")
    add_equation(doc, "Re_j = 0.5 / RP_j,  when risk is active, AP_j = 0, and RP_j > 0.")
    add_equation(doc, "Re_j = 0,  when risk is active without recovery.")
    add_equation(doc, "Re_j = 1 − (B_t,j + U_t,j) / j,  when no risk is active.")

    for sentence in (
        "The formula was recomputed on 47,546 source-workbook rows with zero mismatches and maximum absolute error equal to zero.",
        "The visible workbook population contains completed non-lost orders.",
        "Lost and horizon-unresolved orders are therefore reported through mandatory companion outcomes rather than silently treated as successful rows.",
        "The policy receives zero intermediate reward for decisions one through seven.",
        "The terminal reward is the official visible ReT at the common score time.",
        "This sparse reward prevents the study from introducing an unvalidated shaping function that changes the construct.",
        "The paper should describe the objective as maximizing official ReT subject to fixed resources and complete ledger disclosure.",
    ):
        add_sentence(doc, sentence)

    add_callout(
        doc,
        "Metric interpretation",
        [
            "ReT is the primary construct requested by Garrido.",
            "ReT alone does not describe which product receives service or how many orders remain unresolved.",
            "Worst-product fill, lost demand, unresolved demand, terminal inventory, resource use, and mass conservation must therefore accompany the primary mean.",
        ],
        fill=AMBER,
    )

    add_heading(doc, "11. The learning context", 1)
    for sentence in (
        "One training episode is one eight-week product-mix campaign after the common warm-up.",
        "The agent observes the current state, selects one of four actions, and advances the DES by one decision week.",
        "The recurrent hidden state carries information across the eight decisions within the episode.",
        "The hidden state is reset between independent training episodes.",
        "Program Q therefore studies within-campaign learned feedback rather than accumulated learning across campaigns.",
        "The learner is Stable-Baselines3-Contrib RecurrentPPO with the MlpLstmPolicy implementation.",
        "The recurrent state has sixty-four hidden units.",
        "The policy and value heads use two fully connected layers with sixty-four units each.",
        "The learning rate is 0.0003.",
        "The rollout length is 512 steps.",
        "The minibatch size is sixty-four.",
        "The discount factor is 0.99.",
        "The generalized-advantage parameter is 0.95.",
        "The clipping range is 0.2.",
        "The entropy coefficient is 0.01.",
        "Each of ten learner seeds executes 200,192 training steps.",
        "Only the final checkpoint is evaluated.",
        "The architecture and checkpoint are not selected by searching the confirmation outcomes.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "12. Why RecurrentPPO is the executed learning baseline", 1)
    for sentence in (
        "PPO is a mature policy-gradient algorithm with a clipped update that limits destructive policy changes during training.",
        "The recurrent extension is appropriate because the latent demand regime is not directly observed.",
        "The LSTM provides a compact memory of the observation and action sequence.",
        "The MLP heads map the recurrent representation to action probabilities and value estimates.",
        "This architecture is conventional enough to serve as a strong and reproducible learning baseline.",
        "Its role is not to prove that an LSTM is the best possible representation.",
        "Its role is to test whether a learned feedback policy can outperform every fixed calendar under an exact static benchmark.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "13. DMLPA and KAN as matched architecture tests", 1)
    for sentence in (
        "DMLPA is an exploratory attention-based feature extractor rather than an executed Program Q comparator.",
        "The faithful DMLPA sequence is a linear layer, a GELU activation, a second linear layer, a Transformer encoder, and selection of the last token.",
        "Its hypothesis is that attention over a stacked history can identify temporal dependencies that a conventional representation may miss.",
        "KAN is an exploratory spline-based feature extractor rather than an executed Program Q comparator.",
        "Its hypothesis is that learned univariate spline transformations can expose nonlinear response shapes and permit direct inspection of learned curves.",
        "A KAN is not automatically more parameter-efficient than an MLP.",
        "A KAN is not automatically interpretable merely because its edges are splines.",
        "The current KAN prototype reports approximately 2,249,000 parameters, and the matched rerun must report the baseline checkpoint parameter count next to it.",
        "Parameter count, training time, action latency, spline attribution, stability, and predictive use must be measured rather than asserted.",
        "Preliminary KAN results from another Track B environment cannot be imported into Program Q.",
        "The number of parameters is a model-complexity descriptor and not a resilience outcome.",
        "Any comparison among RecurrentPPO, DMLPA-PPO, and KAN-PPO must use the same DES, observation, action set, reward, training budget, tapes, seeds, evaluation episodes, and resource ledger.",
        "The architecture comparison must be decided by ReT and companion outcomes rather than by training reward alone.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 6. Architecture comparison protocol.")
    add_table(
        doc,
        ["Architecture", "Representation hypothesis", "Strength", "Required caution"],
        [
            (
                "RecurrentPPO with MLP-LSTM",
                "A recurrent state summarizes partial observation across eight weekly decisions.",
                "Executed, reproducible, and directly tied to the Program Q result.",
                "Recurrence does not itself establish cumulative learning or neural superiority.",
            ),
            (
                "DMLPA-PPO",
                "Attention over a stacked history may highlight nonlocal temporal dependencies.",
                "Provides an explicit temporal-representation ablation.",
                "Current DMLPA evidence comes from other contracts and cannot enter Program Q without a matched rerun.",
            ),
            (
                "KAN-PPO",
                "Spline features may capture nonlinear thresholds and expose inspectable response shapes.",
                "Offers a concrete post-training spline and attribution audit.",
                "Interpretability and parameter savings must be demonstrated, and current Track B results are not Program Q results.",
            ),
        ],
        [1900, 2700, 2200, 2560],
        compact=True,
    )

    add_heading(doc, "14. How learning is identified", 1)
    for sentence in (
        "Training reward alone does not establish learning.",
        "A trained-versus-untrained comparison is useful as an implementation diagnostic but is not the primary scientific test.",
        "The primary learning test compares the trained policy with the maximum mean performance over the complete static frontier.",
        "The static benchmark is not a convenient sample because all 65,536 action calendars are enumerated.",
        "The policy must also produce state-dependent action trajectories rather than collapse to one fixed calendar.",
        "Replacement and trajectory audits test whether the action changes when the observed history changes.",
        "The strongest tested structured controller receives the same deployable history.",
        "The comparison with structured feedback asks whether the gain comes from feedback or from neural representation.",
        "A per-episode oracle that selects the best action after seeing the realized outcome is a clairvoyant upper bound.",
        "The oracle cannot serve as a deployable baseline or as proof of learned value.",
        "The current Program Q evidence shows a large feedback advantage over the exact static frontier and practical equivalence with structured feedback.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 7. Current Program Q result by evaluation cell.")
    add_table(
        doc,
        ["Cell", "RL minus static frontier", "LCB95", "RL minus structured", "Interpretation"],
        [
            ("ρ=.75, s=.90", "+0.07952", "+0.06608", "−0.00159", "Feedback win and structured equivalence."),
            ("ρ=.90, s=.75", "+0.07255", "+0.06233", "−0.00072", "Feedback win and structured equivalence."),
            ("ρ=.90, s=.90", "+0.11724", "+0.10614", "−0.00041", "Feedback win and structured equivalence."),
        ],
        [1650, 2050, 1250, 1950, 2460],
        compact=True,
    )

    for sentence in (
        "All ten learner seeds are positive against the static frontier in each cell.",
        "The simultaneous learner-versus-structured intervals lie inside the frozen practical-equivalence margin of plus or minus 0.01.",
        "The strongest tested structured family is faster than RecurrentPPO on the measured hardware.",
        "The measured median action latencies are 0.081834 milliseconds for the structured family and 0.573479 milliseconds for RecurrentPPO.",
        "These latencies are hardware-specific descriptive values and do not establish a universal computational advantage.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "15. Resource equality and the end-of-horizon artifact", 1)
    for sentence in (
        "Garrido identified a valid risk that a finite-horizon policy may accumulate inventory near the end of an episode.",
        "Program Q limits this behavior by fixing the total production entitlement before any policy is evaluated.",
        "Every policy controls the labels of the same twenty-four batch slots and cannot request extra shifts or extra production.",
        "Every policy also receives the same fixed-clock downstream transport entitlement.",
        "The controller therefore cannot improve ReT by purchasing more scheduled capacity.",
        "The controller can still allocate the final batches to the wrong product and leave product-specific inventory after demand stops.",
        "A common clearance tail prevents the score from being taken immediately after the final decision.",
        "Terminal P_C inventory, terminal P_H inventory, unresolved orders, unresolved quantity, product fill, and actual transport utilization are recorded.",
        "No positive reward is assigned to ending inventory.",
        "A policy that leaves inventory while orders remain unresolved is exposed by the product ledger and service outcomes.",
        "Average shift use is not a decision variable in Program Q because shift capacity is fixed.",
        "A future Pareto analysis should therefore use actual utilization, unresolved demand, product fill, and computational cost rather than an invented shift decision.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 8. Resource and ledger audit.")
    add_table(
        doc,
        ["Audit field", "Frozen rule", "Purpose"],
        [
            ("Controlled production", "24 batch slots and 120,000 ration units for every calendar.", "Prevents a policy from buying ReT with more production."),
            ("Downstream freight", "The same fixed daily entitlement is reserved whether loaded or empty.", "Separates scheduled capacity from realized utilization."),
            ("Demand", "The same event-keyed demand and product-label tape is used across matched policies.", "Creates a causal paired comparison."),
            ("Product ledger", "P_C and P_H inventory, demand, backlog, in-flight quantity, and fill are tracked separately.", "Detects cross-product starvation and terminal misallocation."),
            ("Unresolved and lost", "All generated orders remain in the ledger at the common score time.", "Prevents disappearance of difficult orders."),
            ("Mass balance", "Raw material, production, transport, delivery, and ending inventory must reconcile.", "Detects numerical creation or deletion of units."),
        ],
        [2200, 3700, 3460],
        compact=True,
    )

    add_heading(doc, "16. Sensitive assumptions and validity checks", 1)
    for sentence in (
        "The model's most sensitive assumptions are those that determine what information is available, when actions become physical, and which orders enter the metric.",
        "These assumptions must be stated because small timing differences can change the apparent value of feedback.",
        "The product classes are synthetic and share identical production physics.",
        "The current experiment therefore isolates allocation under demand uncertainty rather than product-specific manufacturing complexity.",
        "Storage is unlimited, vehicle availability is represented by fixed entitlements, and product substitution is disabled.",
        "The primary experiment is risk-off.",
        "The result therefore concerns nonstationary demand composition and not response to active disruptions.",
        "The visible ReT population omits lost and unresolved orders from its mean.",
        "The companion ledger is consequently part of the scientific result rather than an optional appendix.",
        "The Python reconstruction has not passed every held-out queue and recovery-tail validation gate against the thesis workbooks.",
        "The paper must use transcription language and must not claim complete endogenous replication of the thesis experiments.",
    ):
        add_sentence(doc, sentence)

    add_caption(doc, "Table 9. Sensitive modeling decisions.")
    add_table(
        doc,
        ["Decision", "Why it matters", "Current control"],
        [
            ("Product commitment after 24 h", "A shorter delay would give the controller more immediate authority.", "Activation delay is frozen and relabeling is forbidden."),
            ("Equal-time ordering", "A lot arriving at the release timestamp could otherwise receive an artificial same-time benefit.", "Half-open event semantics and scheduler-invariance tests are enforced."),
            ("Visible ReT population", "Completed-only rows can conceal unresolved demand.", "Unresolved, lost, fill, and inventory are mandatory companion outcomes."),
            ("Fixed HMM parameters", "True evaluation-cell parameters would leak privileged information.", "One fixed deployable belief model is used in every cell."),
            ("Risk-off primary", "Active risks could confound product-mix feedback with disruption exposure.", "Risk claims are explicitly excluded from Program Q."),
            ("Common score time", "Different clearance time would alter unresolved orders and apparent resilience.", "All policies use the same 1,344-hour clearance tail."),
            ("Static-frontier definition", "A sampled static baseline can make feedback look stronger than it is.", "All 4⁸ calendars are enumerated."),
            ("Structured controller", "A weak classical baseline can manufacture a neural premium.", "The strongest tested structured family is reselected inside the bootstrap."),
        ],
        [2300, 3430, 3630],
        compact=True,
    )

    add_heading(doc, "17. Supported and unsupported manuscript statements", 1)
    add_status_table(
        doc,
        [
            (
                "SUPPORTED",
                "The operational Excel ReT formula is reproduced exactly on 47,546 source rows.",
                "State the row count, zero mismatches, and the visible-population caveat together.",
            ),
            (
                "SUPPORTED",
                "RecurrentPPO beats every policy in the complete Program Q open-loop frontier.",
                "Qualify the claim by the disclosed two-product, risk-off, eight-week contract.",
            ),
            (
                "SUPPORTED",
                "The value of feedback is robust, while the neural premium is not identified.",
                "Report practical equivalence with the strongest tested structured family.",
            ),
            (
                "CONDITIONAL",
                "KAN or DMLPA provides a better representation.",
                "State this only after a matched Program Q run with paired seeds, tapes, reward, budget, and resource ledgers.",
            ),
            (
                "CONDITIONAL",
                "KAN provides interpretable decision logic.",
                "Require stable spline attribution, policy-level action analysis, and a link between the explanation and outcome differences.",
            ),
            (
                "NOT SUPPORTED",
                "The Python DES faithfully reproduces all Garrido thesis experiments.",
                "Use high-fidelity transcription with causal repairs and explicit held-out validation gaps.",
            ),
            (
                "NOT SUPPORTED",
                "RecurrentPPO is safer, faster, or superior to structured feedback.",
                "Report structured equivalence, failed worst-product non-inferiority, and hardware-specific latency.",
            ),
            (
                "NOT SUPPORTED",
                "Program Q cures the cross-campaign Alzheimer effect.",
                "Reserve accumulated learning for the separate retained-learning study.",
            ),
        ],
    )

    add_heading(doc, "18. Manuscript-ready methods narrative", 1)
    for sentence in (
        "We reconstructed the military food supply-chain DES reported by Garrido-Rios as a sequence of thirteen procurement, production, storage, transport, and demand operations.",
        "The reconstruction transcribes the source topology, operating cadences, risk tables, and operational ReT formula while preserving an explicit ledger for orders, products, resources, and mass.",
        "We then introduced a researcher-defined two-product extension to isolate a weekly allocation decision under partially observed changes in demand composition.",
        "The extension contains two nonfungible product classes with identical physical requirements and a shared production line.",
        "The controller assigns the next three fixed 5,000-unit batch targets each week without changing total production or scheduled transport capacity.",
        "The treatment horizon contains eight weekly decisions and therefore defines 4⁸ or 65,536 possible open-loop calendars.",
        "We enumerate this entire static frontier rather than sampling a subset of convenient baselines.",
        "Demand quantities follow the thesis range of 2,400 to 2,600 rations, while product identity follows a latent two-state Markov regime that is never disclosed to the controller.",
        "The policy observes twenty-one normalized operational fields covering inventory, pipeline, backlog, in-flight quantities, a fixed deployable belief, the previous action, and the public episode phase.",
        "We train a RecurrentPPO MLP-LSTM policy because the current snapshot does not fully reveal the latent demand regime.",
        "The reward is zero at intermediate decisions and equals the official terminal ReT at the common score time.",
        "We compare the trained policy with the complete static frontier and with the strongest tested structured feedback family that receives the same deployable information.",
        "This comparison separates the value of feedback from any value attributable specifically to a neural controller.",
        "We report unresolved demand, lost demand, product-level fill, terminal product inventory, actual utilization, and mass conservation alongside ReT.",
        "The resulting evidence supports learned feedback over every open-loop calendar and practical equivalence with structured feedback, but it does not support neural superiority or worst-product safety.",
    ):
        add_sentence(doc, sentence)

    add_heading(doc, "Appendix A. Coauthor integration checklist", 1)
    add_callout(
        doc,
        "Remove this appendix before submission",
        [
            "This appendix is an editorial control surface for the shared working document.",
            "It is not part of the proposed journal manuscript.",
        ],
        fill=LIGHT_GRAY,
    )
    checklist = [
        ("Replace v0 Section 3.2", "Use Sections 1–10 of this insert as the detailed DES and learning-environment source."),
        ("Replace v0 Section 3.3", "Remove the obsolete five-dimensional action space, fifteen-dimensional observation, and shaped reward."),
        ("Baseline wording", "Call Garrido the static reference and call the 65,536-calendar frontier the matched static benchmark."),
        ("Architecture wording", "Use RecurrentPPO with MLP-LSTM for Program Q and label DMLPA and KAN as matched sidecars pending rerun."),
        ("Metric wording", "Keep ReT official and disclose visible rows, unresolved, lost, product fill, resources, and mass."),
        ("End-horizon wording", "Explain fixed production rights, the clearance tail, terminal product inventory, and unresolved demand."),
        ("Fidelity wording", "Use transcription with causal repairs and held-out validation gaps rather than faithful reproduction."),
        ("Results wording", "State feedback value, structured equivalence, and absent worst-product safety without claiming neural premium."),
        ("Literature wording", "Do not claim first-of-kind novelty until the current review is complete."),
        ("KAN wording", "Do not claim parameter savings or interpretability without a matched Program Q result and a policy-level explanation audit."),
    ]
    add_table(
        doc,
        ["Editorial action", "Required change"],
        checklist,
        [2500, 6860],
        compact=True,
    )

    add_heading(doc, "Appendix B. Internal source traceability", 1)
    for sentence in (
        "The following files were used to control technical accuracy during drafting.",
        "Repository paths are internal provenance and can be moved to a reproducibility supplement before submission.",
    ):
        add_sentence(doc, sentence)
    add_table(
        doc,
        ["Source", "Controlled fact"],
        [
            ("research/paper2_exhaustive_search/source_reconstruction.md", "Garrido source map, Op1–Op13 semantics, assumptions, risk definitions, and fidelity boundary."),
            ("contracts/program_o_full_des_hpi_translation_v1.json", "Two-product physics, action timing, demand tape, warm-up, clearance, resources, and same-time rules."),
            ("contracts/program_o_ret_only_learner_v1.json", "Observation contract, RecurrentPPO hyperparameters, comparator ladder, and claim boundaries."),
            ("supply_chain/program_o_ret_env.py", "Twenty-one-dimensional normalization and episode reward implementation."),
            ("supply_chain/ret_thesis.py", "Operational ReT branch logic and ledger semantics."),
            ("papers/submission_a_program_q/source_of_truth.json", "Executed Program Q design, numeric results, guardrails, latency, supported claims, and prohibited claims."),
            ("docs/GARRIDO_FIDELITY_AUDIT_2026-07-10.md", "Static transcription status, causal repairs, held-out validation failures, and allowed language."),
            ("docs/REAL_KAN_10SEED_EXTENSION_VERDICT_2026-07-03.md", "Track B KAN sidecar status and the prohibition on importing it into Program Q."),
        ],
        [3950, 5410],
        compact=True,
    )

    add_sentence(
        doc,
        "End of manuscript-ready DES and learning-environment insert.",
        style="Small Note",
    )
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_reviewed_flow_figure()
    doc = build_document()
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
